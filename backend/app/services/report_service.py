"""Report generation service — Big 4 style professional reports with screenshots."""
from datetime import datetime
from io import BytesIO
from pathlib import Path
import base64
import hashlib
import json
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from app.services.compliance_mapping import get_compliance_mapping
from app.core.config import get_settings

try:
    from fpdf import FPDF
except ImportError:
    FPDF = None


def _severity_order(s: str) -> int:
    order = {"critical": 0, "high": 1, "medium": 2, "low": 3, "info": 4}
    return order.get(s.lower(), 5)


def _risk_score(project: dict, findings: list) -> float:
    weights = {"critical": 25, "high": 15, "medium": 8, "low": 3, "info": 1}
    total = sum(weights.get(f.get("severity", "").lower(), 0) for f in findings)
    base = 10 if project.get("failed_count", 0) > 0 else 0
    return min(100, total + base)


def _coverage_pct(project: dict) -> float:
    total = project.get("total_test_cases") or 0
    if total == 0:
        return 0.0
    tested = project.get("tested_count") or 0
    return round((tested / total) * 100, 1)


def _risk_level(score: float) -> str:
    if score >= 75:
        return "Critical"
    if score >= 50:
        return "High"
    if score >= 25:
        return "Medium"
    return "Low"


def _get_evidence_file_path(project_id: str, evidence_item: dict) -> Path | None:
    """Resolve evidence item to file path. evidence_item: {filename, url}."""
    settings = get_settings()
    base = Path(settings.uploads_path) / str(project_id)
    url = evidence_item.get("url") or ""
    match = re.search(r"/evidence/([^/]+)$", url)
    if match:
        fpath = base / match.group(1)
        if fpath.exists():
            return fpath
    return None


def _evidence_to_base64(project_id: str, evidence_item: dict) -> str | None:
    """Read evidence file and return base64 data URL for images."""
    fpath = _get_evidence_file_path(project_id, evidence_item)
    if not fpath or not fpath.exists():
        return None
    ext = fpath.suffix.lower()
    if ext not in {".png", ".jpg", ".jpeg", ".gif", ".webp"}:
        return None
    try:
        data = fpath.read_bytes()
        b64 = base64.b64encode(data).decode()
        mime = {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg", "gif": "image/gif", "webp": "image/webp"}.get(ext[1:], "image/png")
        return f"data:{mime};base64,{b64}"
    except Exception:
        return None


def build_report_data(project: dict, findings: list, phases: list, project_id: str = "", organization: dict | None = None) -> dict:
    """Build structured report data with evidence."""
    risk = _risk_score(project, findings)
    coverage = _coverage_pct(project)
    severity_counts = {}
    for f in findings:
        s = f.get("severity", "info").lower()
        severity_counts[s] = severity_counts.get(s, 0) + 1
    owasp_map = {}
    cwe_map = {}
    for f in findings:
        o = f.get("owasp_category") or "Other"
        owasp_map[o] = owasp_map.get(o, 0) + 1
        c = f.get("cwe_id") or "N/A"
        cwe_map[c] = cwe_map.get(c, 0) + 1
    sorted_findings = sorted(findings, key=lambda x: _severity_order(x.get("severity", "")))
    for f in sorted_findings:
        f["compliance"] = get_compliance_mapping(cwe_id=f.get("cwe_id"), owasp_category=f.get("owasp_category"))
    org = organization or {"name": "AppSecD", "logo_base64": None, "brand_color": "#2563eb"}
    ai = project.get("ai_report_content") or {}
    return {
        "project": project,
        "organization": org,
        "ai_report_content": ai,
        "findings": sorted_findings,
        "phases": phases,
        "project_id": project_id,
        "risk_score": risk,
        "risk_level": _risk_level(risk),
        "coverage_pct": coverage,
        "severity_distribution": severity_counts,
        "owasp_mapping": owasp_map,
        "cwe_mapping": cwe_map,
        "compliance_summary": {"OWASP Top 10": True, "CWE Top 25": True, "MITRE ATT&CK": True, "ISO 27001 Annex A": True, "NIST 800-53": True},
        "generated_at": datetime.utcnow().isoformat() + "Z",
    }


def _html_escape(s: str) -> str:
    if not s:
        return ""
    return str(s).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")


def generate_html(data: dict) -> str:
    """Generate Big 4 style professional HTML report with screenshots."""
    p = data["project"]
    org = data.get("organization") or {"name": "AppSecD", "logo_base64": None, "brand_color": "#2563eb"}
    org_name = org.get("name") or "AppSecD"
    org_logo = org.get("logo_base64")
    org_color = org.get("brand_color") or "#2563eb"
    findings = data["findings"]
    risk = data["risk_score"]
    risk_level = data.get("risk_level", "Medium")
    cov = data["coverage_pct"]
    sev = data["severity_distribution"]
    owasp = data["owasp_mapping"]
    cwe = data["cwe_mapping"]
    gen = data["generated_at"]
    project_id = data.get("project_id", "")

    # Severity badge HTML helper
    def _sev_badge(severity: str) -> str:
        s = severity.lower()
        return f'<span class="sev-badge sev-badge-{s}">{severity}</span>'

    # AI report content section (when AI features utilized)
    def _ai_exec_section(d: dict) -> str:
        ai = d.get("ai_report_content") or {}
        if not ai:
            return ""
        out = []
        exec_sum = ai.get("executive_summary") or ai.get("ai_summary")
        if exec_sum:
            out.append(f'<div class="ai-exec-block"><div class="ai-exec-label">AI-Generated Executive Summary</div><p class="ai-exec-text">{_html_escape(str(exec_sum))}</p></div>')
        tech = ai.get("technical_summary")
        if tech:
            out.append(f'<div class="ai-exec-block"><div class="ai-exec-label">Technical Summary</div><p class="ai-exec-text">{_html_escape(str(tech))}</p></div>')
        strat = ai.get("strategic_recommendations")
        if strat:
            if isinstance(strat, list):
                strat = "\n".join(f"• {s}" for s in strat)
            out.append(f'<div class="ai-exec-block"><div class="ai-exec-label">Strategic Recommendations</div><p class="ai-exec-text" style="white-space:pre-wrap;">{_html_escape(str(strat))}</p></div>')
        if not out:
            return ""
        return '<div class="ai-exec-summary">' + "".join(out) + "</div>"

    # Build severity distribution rows
    sev_rows = "".join(
        f'<tr><td>{_sev_badge(k.capitalize())}</td><td class="count-cell">{v}</td></tr>'
        for k, v in sorted(sev.items(), key=lambda x: _severity_order(x[0]))
    )
    owasp_rows = "".join(f'<tr><td>{_html_escape(k)}</td><td class="count-cell">{v}</td></tr>' for k, v in sorted(owasp.items()))
    cwe_rows = "".join(f'<tr><td><code>{_html_escape(k)}</code></td><td class="count-cell">{v}</td></tr>' for k, v in sorted(cwe.items()))

    # Build findings summary table rows
    find_rows = ""
    for i, f in enumerate(findings, 1):
        find_rows += f"""
        <tr>
            <td class="count-cell">{i}</td>
            <td><a href="#finding-{i}" class="finding-link">{_html_escape(f.get('title', ''))}</a></td>
            <td>{_sev_badge(f.get('severity',''))}</td>
            <td>{_html_escape(f.get('owasp_category') or '-')}</td>
            <td><code>{_html_escape(f.get('cwe_id') or '-')}</code></td>
            <td class="url-cell">{_html_escape(f.get('affected_url') or '-')}</td>
        </tr>
        """

    # Executive summary bullets
    exec_bullets = []
    if findings:
        critical_high = sum(1 for x in findings if x.get("severity", "").lower() in ("critical", "high"))
        if critical_high:
            exec_bullets.append(f"<li><strong>{critical_high} Critical/High severity finding(s)</strong> require immediate remediation.</li>")
        exec_bullets.append(f"<li>Test coverage: <strong>{cov}%</strong> ({p.get('tested_count',0)} of {p.get('total_test_cases',0)} test cases executed).</li>")
        exec_bullets.append(f"<li>Findings mapped to OWASP Top 10, CWE Top 25, MITRE ATT&amp;CK, ISO 27001, and NIST 800-53.</li>")
        exec_bullets.append("<li>Recommendations provided for each finding to support remediation planning.</li>")
    else:
        exec_bullets.append("<li>No security findings identified during this assessment.</li>")
        exec_bullets.append(f"<li>Test coverage: <strong>{cov}%</strong>.</li>")
    exec_bullets_html = "".join(exec_bullets)

    scope = _html_escape(p.get("testing_scope") or "In-scope: Application under test. Out-of-scope: Third-party systems.")
    methodology = "Black-box" if p.get("testing_type") == "black_box" else "Grey-box" if p.get("testing_type") == "grey_box" else "White-box"
    methodology += f" testing in {p.get('environment', 'staging')} environment. Manual and semi-automated techniques per OWASP Testing Guide."

    sev_colors = ["#dc2626", "#ea580c", "#ca8a04", "#16a34a", "#2563eb"]

    # Risk matrix data
    risk_level_class = risk_level.lower()

    # Build KPI cards
    total_findings = len(findings)
    critical_count = sev.get("critical", 0)
    high_count = sev.get("high", 0)

    # Build compliance framework rows for the compliance table
    compliance_rows = ""
    frameworks = [
        ("OWASP Top 10 (2021)", "Web Application Security Standard", "Mapped"),
        ("CWE Top 25", "Common Weakness Enumeration", "Mapped"),
        ("MITRE ATT&amp;CK", "Adversarial Tactics and Techniques", "Mapped"),
        ("ISO 27001 Annex A", "Information Security Controls", "Mapped"),
        ("NIST 800-53", "Security and Privacy Controls", "Mapped"),
    ]
    for fw_name, fw_desc, fw_status in frameworks:
        compliance_rows += f'<tr><td><strong>{fw_name}</strong></td><td>{fw_desc}</td><td><span class="compliance-mapped">{fw_status}</span></td></tr>'

    # Build findings trend by date
    from collections import defaultdict
    trend_by_date = defaultdict(int)
    for f in findings:
        ca = f.get("created_at")
        if ca and len(ca) >= 10:
            trend_by_date[ca[:10]] += 1
    trend_dates = sorted(trend_by_date.keys())
    trend_counts = [trend_by_date[d] for d in trend_dates]
    if not trend_dates and findings:
        trend_dates = [gen[:10]]
        trend_counts = [len(findings)]

    # Build CWE detail rows for the CWE mapping table
    cwe_detail_rows = ""
    for f in findings:
        cwe_id = f.get("cwe_id") or "-"
        comp = f.get("compliance") or {}
        cwe_detail_rows += f"""<tr>
            <td><code>{_html_escape(cwe_id)}</code></td>
            <td>{_html_escape(f.get('title', ''))}</td>
            <td>{_sev_badge(f.get('severity', ''))}</td>
            <td><code>{_html_escape(comp.get('mitre_attack') or '-')}</code></td>
            <td>{_html_escape(comp.get('iso_27001') or '-')}</td>
            <td>{_html_escape(comp.get('nist_800_53') or '-')}</td>
        </tr>"""

    toc_items = [
        ("1. Executive Summary", "exec-summary"),
        ("2. Scope &amp; Methodology", "scope"),
        ("3. Risk Rating Methodology", "methodology"),
        ("4. Risk Matrix", "risk-matrix"),
        ("5. Charts &amp; Analytics", "charts"),
        ("6. Summary Statistics", "stats"),
        ("7. OWASP Top 10 Mapping", "owasp"),
        ("8. CWE Mapping", "cwe"),
        ("9. Compliance Frameworks", "compliance"),
        ("10. Findings Summary", "findings-table"),
        ("11. Finding Details with Evidence", "finding-details"),
        ("12. Remediation Timeline", "timeline"),
    ]
    toc_html = "".join(f'<li><a href="#{tid}" class="toc-link">{txt}</a></li>' for txt, tid in toc_items)

    # Timeline items for findings
    timeline_html = ""
    for i, f in enumerate(findings, 1):
        sev_lower = f.get("severity", "").lower()
        if sev_lower == "critical":
            timeline_label = "Immediate (0-7 days)"
        elif sev_lower == "high":
            timeline_label = "Urgent (7-30 days)"
        elif sev_lower == "medium":
            timeline_label = "Planned (30-90 days)"
        else:
            timeline_label = "Backlog (Next release)"
        timeline_html += f"""
        <div class="timeline-item timeline-{sev_lower}">
            <div class="timeline-marker"></div>
            <div class="timeline-content">
                <div class="timeline-header">
                    <span class="timeline-title">Finding #{i}: {_html_escape(f.get('title',''))}</span>
                    {_sev_badge(f.get('severity',''))}
                </div>
                <div class="timeline-target">{timeline_label}</div>
            </div>
        </div>"""

    html = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Security Assessment Report - {_html_escape(p.get('application_name', ''))}</title>
    <style>
        /* ============================================================
           CSS Reset & Base
           ============================================================ */
        @page {{ margin: 2cm; }}
        *, *::before, *::after {{ box-sizing: border-box; }}
        :root {{
            --primary: #1e3a5f;
            --primary-light: #2563eb;
            --primary-dark: #0f2440;
            --accent: #0ea5e9;
            --bg-body: #f0f4f8;
            --bg-white: #ffffff;
            --text-primary: #1a202c;
            --text-secondary: #4a5568;
            --text-muted: #718096;
            --border: #e2e8f0;
            --border-light: #edf2f7;
            --shadow-sm: 0 1px 3px rgba(0,0,0,0.08);
            --shadow-md: 0 4px 12px rgba(0,0,0,0.1);
            --shadow-lg: 0 10px 30px rgba(0,0,0,0.12);
            --radius: 12px;
            --radius-sm: 8px;
            --radius-xs: 4px;
            --gradient-brand: linear-gradient(135deg, #1e3a5f 0%, #2563eb 50%, #0ea5e9 100%);
            --gradient-dark: linear-gradient(135deg, #0f2440 0%, #1e3a5f 100%);
            --sev-critical: #dc2626;
            --sev-high: #ea580c;
            --sev-medium: #d97706;
            --sev-low: #16a34a;
            --sev-info: #2563eb;
            --sev-critical-bg: #fef2f2;
            --sev-high-bg: #fff7ed;
            --sev-medium-bg: #fffbeb;
            --sev-low-bg: #f0fdf4;
            --sev-info-bg: #eff6ff;
        }}
        body {{
            font-family: 'Segoe UI', system-ui, -apple-system, BlinkMacSystemFont, 'Helvetica Neue', Arial, sans-serif;
            margin: 0; padding: 0;
            color: var(--text-primary);
            line-height: 1.7;
            background: var(--bg-body);
            -webkit-font-smoothing: antialiased;
        }}
        .report-container {{
            max-width: 960px;
            margin: 0 auto;
            padding: 0 1.5rem;
        }}

        /* ============================================================
           Cover Page
           ============================================================ */
        .cover-page {{
            background: var(--gradient-brand);
            color: white;
            padding: 4rem 3rem;
            text-align: center;
            position: relative;
            overflow: hidden;
        }}
        .cover-page::before {{
            content: '';
            position: absolute; top: 0; left: 0; right: 0; bottom: 0;
            background: url("data:image/svg+xml,%3Csvg width='60' height='60' viewBox='0 0 60 60' xmlns='http://www.w3.org/2000/svg'%3E%3Cg fill='none' fill-rule='evenodd'%3E%3Cg fill='%23ffffff' fill-opacity='0.05'%3E%3Cpath d='M36 34v-4h-2v4h-4v2h4v4h2v-4h4v-2h-4zm0-30V0h-2v4h-4v2h4v4h2V6h4V4h-4zM6 34v-4H4v4H0v2h4v4h2v-4h4v-2H6zM6 4V0H4v4H0v2h4v4h2V6h4V4H6z'/%3E%3C/g%3E%3C/g%3E%3C/svg%3E");
            opacity: 0.3;
        }}
        .cover-page * {{ position: relative; z-index: 1; }}
        .cover-logo {{
            font-size: 1rem;
            letter-spacing: 6px;
            text-transform: uppercase;
            opacity: 0.85;
            margin-bottom: 1.5rem;
            font-weight: 300;
        }}
        .cover-title {{
            font-size: 2.5rem;
            font-weight: 700;
            margin: 0 0 0.5rem 0;
            letter-spacing: -0.5px;
            text-shadow: 0 2px 4px rgba(0,0,0,0.2);
        }}
        .cover-subtitle {{
            font-size: 1.25rem;
            opacity: 0.9;
            margin: 0.75rem 0 2rem 0;
            font-weight: 300;
        }}
        .cover-meta {{
            display: flex;
            justify-content: center;
            gap: 2rem;
            flex-wrap: wrap;
            margin-top: 2rem;
        }}
        .cover-meta-item {{
            background: rgba(255,255,255,0.15);
            backdrop-filter: blur(10px);
            padding: 0.75rem 1.5rem;
            border-radius: var(--radius-sm);
            border: 1px solid rgba(255,255,255,0.2);
            font-size: 0.85rem;
        }}
        .cover-meta-label {{
            opacity: 0.75;
            font-size: 0.7rem;
            text-transform: uppercase;
            letter-spacing: 1px;
            display: block;
            margin-bottom: 2px;
        }}
        .cover-classification {{
            display: inline-block;
            margin-top: 2rem;
            padding: 0.4rem 1.5rem;
            border: 2px solid rgba(255,255,255,0.5);
            border-radius: var(--radius-xs);
            font-size: 0.8rem;
            text-transform: uppercase;
            letter-spacing: 3px;
            font-weight: 600;
        }}

        /* ============================================================
           Typography
           ============================================================ */
        .report-body {{ padding: 2rem 0 4rem 0; }}
        h1, h2, h3, h4, h5, h6 {{ margin-top: 0; font-weight: 700; }}
        h2.section-heading {{
            color: var(--primary);
            font-size: 1.5rem;
            margin: 3rem 0 1.25rem 0;
            padding-bottom: 0.75rem;
            border-bottom: 3px solid var(--primary-light);
            display: flex;
            align-items: center;
            gap: 0.75rem;
        }}
        h2.section-heading .section-number {{
            background: var(--gradient-brand);
            color: white;
            width: 36px; height: 36px;
            border-radius: 50%;
            display: inline-flex;
            align-items: center;
            justify-content: center;
            font-size: 0.85rem;
            flex-shrink: 0;
        }}
        h3 {{ color: var(--text-primary); margin-top: 1.5rem; font-size: 1.1rem; }}
        p {{ margin: 0.75rem 0; }}
        code {{
            background: #edf2f7;
            padding: 0.15em 0.4em;
            border-radius: 3px;
            font-size: 0.88em;
            font-family: 'Fira Code', 'Cascadia Code', 'SF Mono', Consolas, monospace;
        }}

        /* ============================================================
           Cards
           ============================================================ */
        .card {{
            background: var(--bg-white);
            border-radius: var(--radius);
            padding: 1.75rem;
            margin: 1.25rem 0;
            box-shadow: var(--shadow-sm);
            border: 1px solid var(--border-light);
            transition: box-shadow 0.2s ease;
        }}
        .card:hover {{ box-shadow: var(--shadow-md); }}
        .card-header {{
            font-weight: 700;
            font-size: 1.05rem;
            color: var(--primary);
            margin-bottom: 0.75rem;
        }}

        /* ============================================================
           KPI Cards
           ============================================================ */
        .kpi-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
            gap: 1rem;
            margin: 1.5rem 0;
        }}
        .kpi-card {{
            background: var(--bg-white);
            border-radius: var(--radius);
            padding: 1.5rem;
            text-align: center;
            box-shadow: var(--shadow-sm);
            border: 1px solid var(--border-light);
            position: relative;
            overflow: hidden;
        }}
        .kpi-card::before {{
            content: '';
            position: absolute;
            top: 0; left: 0; right: 0;
            height: 4px;
        }}
        .kpi-card.kpi-risk::before {{ background: var(--sev-{risk_level_class}); }}
        .kpi-card.kpi-findings::before {{ background: var(--primary-light); }}
        .kpi-card.kpi-coverage::before {{ background: var(--accent); }}
        .kpi-card.kpi-critical::before {{ background: var(--sev-critical); }}
        .kpi-value {{
            font-size: 2.25rem;
            font-weight: 800;
            line-height: 1.2;
        }}
        .kpi-value.risk-val {{ color: var(--sev-{risk_level_class}); }}
        .kpi-value.findings-val {{ color: var(--primary-light); }}
        .kpi-value.coverage-val {{ color: var(--accent); }}
        .kpi-value.critical-val {{ color: var(--sev-critical); }}
        .kpi-label {{
            font-size: 0.8rem;
            text-transform: uppercase;
            letter-spacing: 1px;
            color: var(--text-muted);
            margin-top: 0.25rem;
            font-weight: 600;
        }}

        /* ============================================================
           Document Control
           ============================================================ */
        .doc-control {{
            background: var(--bg-white);
            padding: 1.25rem 1.5rem;
            border-radius: var(--radius-sm);
            font-size: 0.85rem;
            margin: 1.5rem 0;
            border: 1px solid var(--border);
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
            gap: 0.75rem;
        }}
        .doc-control-item {{ }}
        .doc-control-label {{
            font-size: 0.7rem;
            text-transform: uppercase;
            letter-spacing: 1px;
            color: var(--text-muted);
            font-weight: 600;
        }}
        .doc-control-value {{
            color: var(--text-primary);
            font-weight: 500;
        }}

        /* ============================================================
           Table of Contents
           ============================================================ */
        .toc {{
            background: var(--bg-white);
            padding: 1.75rem 2rem;
            border-radius: var(--radius);
            margin: 1.5rem 0;
            border: 1px solid var(--border-light);
            box-shadow: var(--shadow-sm);
        }}
        .toc-title {{
            font-weight: 700;
            font-size: 1.1rem;
            color: var(--primary);
            margin-bottom: 0.75rem;
        }}
        .toc ol {{
            margin: 0;
            padding-left: 1.25rem;
            list-style: none;
            counter-reset: toc-counter;
        }}
        .toc li {{
            counter-increment: toc-counter;
            padding: 0.4rem 0;
            border-bottom: 1px solid var(--border-light);
        }}
        .toc li:last-child {{ border-bottom: none; }}
        .toc-link {{
            color: var(--primary-light);
            text-decoration: none;
            font-weight: 500;
            transition: color 0.15s;
        }}
        .toc-link:hover {{
            color: var(--primary);
            text-decoration: underline;
        }}

        /* ============================================================
           Executive Summary
           ============================================================ */
        .exec-summary {{
            background: linear-gradient(135deg, #eff6ff 0%, #dbeafe 50%, #e0f2fe 100%);
            padding: 2rem;
            border-radius: var(--radius);
            margin: 1.5rem 0;
            border-left: 5px solid var(--primary-light);
            box-shadow: var(--shadow-sm);
        }}
        .exec-summary p {{ margin: 0.5rem 0; }}
        .exec-summary ul {{ margin: 0.75rem 0; padding-left: 1.5rem; }}
        .exec-summary li {{ margin: 0.35rem 0; }}
        .ai-exec-summary {{ margin-bottom: 1.5rem; }}
        .ai-exec-block {{ background: linear-gradient(135deg, #eff6ff 0%, #e0f2fe 100%); padding: 1.25rem; border-radius: var(--radius-sm); margin-bottom: 1rem; border-left: 4px solid var(--primary-light); }}
        .ai-exec-label {{ font-size: 0.75rem; font-weight: 700; text-transform: uppercase; letter-spacing: 1px; color: var(--primary); margin-bottom: 0.5rem; }}
        .ai-exec-text {{ margin: 0; font-size: 0.95rem; line-height: 1.65; color: var(--text-secondary); }}

        /* ============================================================
           Risk Badge
           ============================================================ */
        .risk-badge {{
            display: inline-block;
            padding: 0.35rem 1rem;
            border-radius: 6px;
            font-weight: 700;
            font-size: 1rem;
            letter-spacing: 0.5px;
        }}
        .risk-critical {{ background: var(--sev-critical); color: white; }}
        .risk-high {{ background: var(--sev-high); color: white; }}
        .risk-medium {{ background: var(--sev-medium); color: white; }}
        .risk-low {{ background: var(--sev-low); color: white; }}

        /* ============================================================
           Severity Badges (inline)
           ============================================================ */
        .sev-badge {{
            display: inline-block;
            padding: 0.2rem 0.65rem;
            border-radius: 20px;
            font-weight: 600;
            font-size: 0.78rem;
            text-transform: uppercase;
            letter-spacing: 0.5px;
            white-space: nowrap;
        }}
        .sev-badge-critical {{ background: var(--sev-critical); color: white; }}
        .sev-badge-high {{ background: var(--sev-high); color: white; }}
        .sev-badge-medium {{ background: var(--sev-medium); color: white; }}
        .sev-badge-low {{ background: var(--sev-low); color: white; }}
        .sev-badge-info {{ background: var(--sev-info); color: white; }}

        /* ============================================================
           Risk Score Gauge
           ============================================================ */
        .risk-gauge-container {{
            display: flex;
            align-items: center;
            gap: 1.5rem;
            margin: 1rem 0;
        }}
        .risk-gauge {{
            width: 100%;
            max-width: 400px;
            height: 24px;
            background: #e2e8f0;
            border-radius: 12px;
            overflow: hidden;
            position: relative;
        }}
        .risk-gauge-fill {{
            height: 100%;
            border-radius: 12px;
            transition: width 0.6s ease;
            background: linear-gradient(90deg, var(--sev-low), var(--sev-medium), var(--sev-high), var(--sev-critical));
            background-size: 400px 100%;
        }}
        .risk-gauge-label {{
            font-weight: 700;
            font-size: 1.1rem;
            white-space: nowrap;
        }}

        /* ============================================================
           Risk Matrix (CSS)
           ============================================================ */
        .risk-matrix {{
            display: grid;
            grid-template-columns: auto repeat(5, 1fr);
            grid-template-rows: auto repeat(5, 1fr);
            gap: 2px;
            max-width: 500px;
            margin: 1.5rem 0;
            font-size: 0.75rem;
        }}
        .risk-matrix-cell {{
            display: flex;
            align-items: center;
            justify-content: center;
            min-height: 44px;
            min-width: 70px;
            border-radius: 4px;
            font-weight: 500;
            text-align: center;
            padding: 4px;
        }}
        .risk-matrix-header {{
            background: var(--primary-dark);
            color: white;
            font-weight: 700;
            font-size: 0.7rem;
            text-transform: uppercase;
            letter-spacing: 0.5px;
        }}
        .risk-matrix-label {{
            background: var(--primary);
            color: white;
            font-weight: 600;
            font-size: 0.72rem;
        }}
        .rm-critical {{ background: #dc2626; color: white; }}
        .rm-high {{ background: #ea580c; color: white; }}
        .rm-medium {{ background: #d97706; color: white; }}
        .rm-low {{ background: #16a34a; color: white; }}
        .rm-info {{ background: #93c5fd; color: #1e3a5f; }}
        .risk-matrix-corner {{
            background: var(--primary-dark);
            color: white;
            font-size: 0.65rem;
            text-transform: uppercase;
            letter-spacing: 0.5px;
        }}

        /* ============================================================
           Tables
           ============================================================ */
        table {{
            border-collapse: separate;
            border-spacing: 0;
            width: 100%;
            margin: 1rem 0;
            font-size: 0.88rem;
            background: var(--bg-white);
            border-radius: var(--radius-sm);
            overflow: hidden;
            box-shadow: var(--shadow-sm);
        }}
        th, td {{
            padding: 0.65rem 1rem;
            text-align: left;
            border-bottom: 1px solid var(--border-light);
        }}
        th {{
            background: var(--primary);
            color: white;
            font-weight: 600;
            font-size: 0.8rem;
            text-transform: uppercase;
            letter-spacing: 0.5px;
        }}
        tr:nth-child(even) td {{ background: #f7fafc; }}
        tr:hover td {{ background: #edf2f7; }}
        tr:last-child td {{ border-bottom: none; }}
        .count-cell {{ text-align: center; font-weight: 600; }}
        .url-cell {{
            max-width: 200px;
            overflow: hidden;
            text-overflow: ellipsis;
            white-space: nowrap;
            font-family: 'Fira Code', monospace;
            font-size: 0.82rem;
        }}
        .finding-link {{
            color: var(--primary-light);
            text-decoration: none;
            font-weight: 600;
        }}
        .finding-link:hover {{ text-decoration: underline; }}

        /* ============================================================
           Compliance Mapped Badge
           ============================================================ */
        .compliance-mapped {{
            display: inline-block;
            padding: 0.15rem 0.6rem;
            background: #dcfce7;
            color: #166534;
            border-radius: 20px;
            font-size: 0.78rem;
            font-weight: 600;
        }}

        /* ============================================================
           Finding Detail Blocks
           ============================================================ */
        .finding-block {{
            margin: 2rem 0;
            background: var(--bg-white);
            border-radius: var(--radius);
            overflow: hidden;
            box-shadow: var(--shadow-sm);
            border: 1px solid var(--border-light);
        }}
        .finding-block-header {{
            padding: 1.25rem 1.75rem;
            display: flex;
            align-items: center;
            justify-content: space-between;
            gap: 1rem;
        }}
        .finding-block-header.sev-bg-critical {{ background: var(--sev-critical-bg); border-bottom: 3px solid var(--sev-critical); }}
        .finding-block-header.sev-bg-high {{ background: var(--sev-high-bg); border-bottom: 3px solid var(--sev-high); }}
        .finding-block-header.sev-bg-medium {{ background: var(--sev-medium-bg); border-bottom: 3px solid var(--sev-medium); }}
        .finding-block-header.sev-bg-low {{ background: var(--sev-low-bg); border-bottom: 3px solid var(--sev-low); }}
        .finding-block-header.sev-bg-info {{ background: var(--sev-info-bg); border-bottom: 3px solid var(--sev-info); }}
        .finding-block-title {{
            font-size: 1.15rem;
            font-weight: 700;
            color: var(--text-primary);
            margin: 0;
        }}
        .finding-block-body {{
            padding: 1.5rem 1.75rem;
        }}
        .finding-field {{
            margin: 1rem 0;
        }}
        .finding-field-label {{
            font-weight: 700;
            font-size: 0.82rem;
            text-transform: uppercase;
            letter-spacing: 0.5px;
            color: var(--text-muted);
            margin-bottom: 0.3rem;
        }}
        .finding-field-value {{
            color: var(--text-primary);
            line-height: 1.7;
        }}
        .finding-meta-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(180px, 1fr));
            gap: 1rem;
            margin: 1rem 0;
            padding: 1rem;
            background: #f7fafc;
            border-radius: var(--radius-sm);
        }}

        /* ============================================================
           Evidence Display
           ============================================================ */
        .evidence-container {{
            margin: 1.25rem 0;
            border: 1px solid var(--border);
            border-radius: var(--radius-sm);
            overflow: hidden;
            background: #f8fafc;
        }}
        .evidence-caption {{
            padding: 0.5rem 1rem;
            background: var(--primary);
            color: white;
            font-size: 0.8rem;
            font-weight: 600;
            text-transform: uppercase;
            letter-spacing: 0.5px;
        }}
        .evidence-img-wrap {{
            padding: 1rem;
            text-align: center;
            background: white;
        }}
        .evidence-img {{
            max-width: 100%;
            height: auto;
            border-radius: var(--radius-xs);
            box-shadow: var(--shadow-md);
        }}
        .evidence-pre {{
            margin: 0.5rem 0;
            padding: 1rem;
            background: #f1f5f9;
            border: 1px solid var(--border);
            border-radius: var(--radius-xs);
            font-size: 11px;
            font-family: 'Consolas', 'Monaco', monospace;
            overflow-x: auto;
            white-space: pre-wrap;
            word-break: break-all;
        }}

        /* ============================================================
           Compliance Tag
           ============================================================ */
        .compliance-tags {{
            display: flex;
            flex-wrap: wrap;
            gap: 0.5rem;
            margin-top: 0.75rem;
        }}
        .compliance-tag {{
            display: inline-block;
            padding: 0.25rem 0.75rem;
            background: #eff6ff;
            color: var(--primary);
            border: 1px solid #bfdbfe;
            border-radius: 20px;
            font-size: 0.78rem;
            font-weight: 500;
        }}

        /* ============================================================
           Remediation Timeline
           ============================================================ */
        .timeline {{
            position: relative;
            padding-left: 2.5rem;
            margin: 1.5rem 0;
        }}
        .timeline::before {{
            content: '';
            position: absolute;
            left: 12px;
            top: 0; bottom: 0;
            width: 3px;
            background: var(--border);
            border-radius: 2px;
        }}
        .timeline-item {{
            position: relative;
            margin-bottom: 1.25rem;
        }}
        .timeline-marker {{
            position: absolute;
            left: -2.5rem;
            top: 0.5rem;
            width: 16px; height: 16px;
            border-radius: 50%;
            border: 3px solid white;
            box-shadow: 0 0 0 2px var(--border);
        }}
        .timeline-critical .timeline-marker {{ background: var(--sev-critical); box-shadow: 0 0 0 2px var(--sev-critical); }}
        .timeline-high .timeline-marker {{ background: var(--sev-high); box-shadow: 0 0 0 2px var(--sev-high); }}
        .timeline-medium .timeline-marker {{ background: var(--sev-medium); box-shadow: 0 0 0 2px var(--sev-medium); }}
        .timeline-low .timeline-marker {{ background: var(--sev-low); box-shadow: 0 0 0 2px var(--sev-low); }}
        .timeline-info .timeline-marker {{ background: var(--sev-info); box-shadow: 0 0 0 2px var(--sev-info); }}
        .timeline-content {{
            background: var(--bg-white);
            padding: 1rem 1.25rem;
            border-radius: var(--radius-sm);
            border: 1px solid var(--border-light);
            box-shadow: var(--shadow-sm);
        }}
        .timeline-header {{
            display: flex;
            align-items: center;
            justify-content: space-between;
            gap: 0.75rem;
            flex-wrap: wrap;
        }}
        .timeline-title {{ font-weight: 600; color: var(--text-primary); }}
        .timeline-target {{
            font-size: 0.82rem;
            color: var(--text-muted);
            margin-top: 0.25rem;
            font-weight: 500;
        }}

        /* ============================================================
           Charts
           ============================================================ */
        .charts-grid {{
            display: flex;
            flex-wrap: wrap;
            gap: 2rem;
            margin: 1.5rem 0;
        }}
        .chart-card {{
            background: var(--bg-white);
            padding: 1.5rem;
            border-radius: var(--radius);
            box-shadow: var(--shadow-sm);
            border: 1px solid var(--border-light);
            flex: 1;
            min-width: 300px;
        }}
        .chart-card-title {{
            font-weight: 700;
            font-size: 0.9rem;
            color: var(--text-secondary);
            text-transform: uppercase;
            letter-spacing: 0.5px;
            margin-bottom: 1rem;
        }}
        .chart-container {{ max-width: 100%; }}

        /* ============================================================
           Footer
           ============================================================ */
        .report-footer {{
            margin-top: 3rem;
            padding: 2rem 0;
            border-top: 3px solid var(--primary);
        }}
        .footer-confidentiality {{
            background: var(--primary-dark);
            color: white;
            padding: 1.5rem 2rem;
            border-radius: var(--radius-sm);
            font-size: 0.82rem;
            line-height: 1.6;
        }}
        .footer-confidentiality strong {{
            display: block;
            font-size: 0.9rem;
            margin-bottom: 0.5rem;
            text-transform: uppercase;
            letter-spacing: 1px;
        }}
        .footer-meta {{
            display: flex;
            justify-content: space-between;
            align-items: center;
            flex-wrap: wrap;
            gap: 1rem;
            margin-top: 1rem;
            font-size: 0.78rem;
            color: var(--text-muted);
        }}

        /* ============================================================
           Print Styles
           ============================================================ */
        @media print {{
            body {{ background: white; -webkit-print-color-adjust: exact; print-color-adjust: exact; }}
            .cover-page {{ break-after: page; }}
            .report-container {{ max-width: 100%; padding: 0; }}
            .card, .kpi-card, .chart-card, table, .finding-block, .toc, .exec-summary {{
                box-shadow: none;
                break-inside: avoid;
            }}
            .finding-block {{ break-inside: avoid; page-break-inside: avoid; }}
            h2.section-heading {{ break-after: avoid; }}
            .charts-grid {{ break-inside: avoid; }}
            tr {{ break-inside: avoid; }}
            .cover-page::before {{ display: none; }}
            .toc {{ page-break-after: always; }}
            .kpi-card::before {{ print-color-adjust: exact; -webkit-print-color-adjust: exact; }}
            .sev-badge, .risk-badge, .compliance-mapped, .compliance-tag {{
                print-color-adjust: exact;
                -webkit-print-color-adjust: exact;
            }}
            a {{ color: inherit; text-decoration: none; }}
            .footer-confidentiality {{
                background: #1a202c !important;
                print-color-adjust: exact;
                -webkit-print-color-adjust: exact;
            }}
        }}
    </style>
    <script src="https://cdn.jsdelivr.net/npm/chart.js@4.4.1/dist/chart.umd.min.js"></script>
</head>
<body>

    <!-- ============================================================
         Cover Page
         ============================================================ -->
    <div class="cover-page">
        <div class="cover-logo-area">
            {f'<img src="{org_logo}" alt="{_html_escape(org_name)}" class="cover-logo-img" style="max-height:60px;max-width:200px;object-fit:contain;" />' if org_logo else f'<div class="cover-logo">{_html_escape(org_name)}</div>'}
        </div>
        <h1 class="cover-title">Application Security<br>Assessment Report</h1>
        <p class="cover-subtitle">{_html_escape(p.get('application_name', '') or p.get('name', ''))}</p>
        <div class="cover-meta">
            <div class="cover-meta-item">
                <span class="cover-meta-label">Version</span>
                {_html_escape(p.get('application_version') or '1.0')}
            </div>
            <div class="cover-meta-item">
                <span class="cover-meta-label">Environment</span>
                {_html_escape(p.get('environment', 'Staging'))}
            </div>
            <div class="cover-meta-item">
                <span class="cover-meta-label">Risk Level</span>
                {risk_level} ({risk}/100)
            </div>
            <div class="cover-meta-item">
                <span class="cover-meta-label">Generated</span>
                {gen[:10]}
            </div>
        </div>
        <div class="cover-classification">{_html_escape(p.get('classification') or 'Confidential')}</div>
    </div>

    <div class="report-container">
        <div class="report-body">

            <!-- Document Control -->
            <div class="doc-control">
                <div class="doc-control-item">
                    <div class="doc-control-label">Report Version</div>
                    <div class="doc-control-value">1.0</div>
                </div>
                <div class="doc-control-item">
                    <div class="doc-control-label">Generated</div>
                    <div class="doc-control-value">{gen}</div>
                </div>
                <div class="doc-control-item">
                    <div class="doc-control-label">Target URL</div>
                    <div class="doc-control-value">{_html_escape(p.get('application_url', ''))}</div>
                </div>
                <div class="doc-control-item">
                    <div class="doc-control-label">Prepared By</div>
                    <div class="doc-control-value">{_html_escape(org_name)}</div>
                </div>
                <div class="doc-control-item">
                    <div class="doc-control-label">Project</div>
                    <div class="doc-control-value">{_html_escape(p.get('name', '') or p.get('application_name', ''))}</div>
                </div>
            </div>

            <!-- Table of Contents -->
            <nav class="toc">
                <div class="toc-title">Table of Contents</div>
                <ol>{toc_html}</ol>
            </nav>

            <!-- KPI Dashboard -->
            <div class="kpi-grid">
                <div class="kpi-card kpi-risk">
                    <div class="kpi-value risk-val">{risk}<span style="font-size:0.6em">/100</span></div>
                    <div class="kpi-label">Risk Score</div>
                </div>
                <div class="kpi-card kpi-findings">
                    <div class="kpi-value findings-val">{total_findings}</div>
                    <div class="kpi-label">Total Findings</div>
                </div>
                <div class="kpi-card kpi-coverage">
                    <div class="kpi-value coverage-val">{cov}%</div>
                    <div class="kpi-label">Test Coverage</div>
                </div>
                <div class="kpi-card kpi-critical">
                    <div class="kpi-value critical-val">{critical_count + high_count}</div>
                    <div class="kpi-label">Critical / High</div>
                </div>
            </div>

            <!-- 1. Executive Summary -->
            <h2 class="section-heading" id="exec-summary"><span class="section-number">1</span> Executive Summary</h2>
            {_ai_exec_section(data)}
            <div class="exec-summary">
                <p>This report presents the findings of a security assessment conducted on <strong>{_html_escape(p.get('application_name', ''))}</strong>.
                The assessment evaluated the application against industry-standard security controls including OWASP Top 10, CWE Top 25, and related frameworks.</p>
                <p><strong>Overall Risk Rating:</strong> <span class="risk-badge risk-{risk_level_class}">{risk_level}</span> ({risk}/100)</p>
                <div class="risk-gauge-container">
                    <div class="risk-gauge">
                        <div class="risk-gauge-fill" style="width: {risk}%;"></div>
                    </div>
                    <span class="risk-gauge-label">{risk}/100</span>
                </div>
                <p><strong>Key Points:</strong></p>
                <ul>{exec_bullets_html}</ul>
                <p><strong>Recommendation:</strong> Address Critical and High severity findings within 30 days. Medium and Low findings should be remediated in subsequent release cycles.</p>
            </div>

            <!-- 2. Scope & Methodology -->
            <h2 class="section-heading" id="scope"><span class="section-number">2</span> Scope &amp; Methodology</h2>
            <div class="card">
                <div class="finding-meta-grid">
                    <div>
                        <div class="finding-field-label">Scope</div>
                        <div class="finding-field-value">{scope}</div>
                    </div>
                    <div>
                        <div class="finding-field-label">Methodology</div>
                        <div class="finding-field-value">{methodology}</div>
                    </div>
                    <div>
                        <div class="finding-field-label">Target Completion</div>
                        <div class="finding-field-value">{_html_escape(p.get('target_completion_date') or 'N/A')}</div>
                    </div>
                    <div>
                        <div class="finding-field-label">Testing Type</div>
                        <div class="finding-field-value">{_html_escape(p.get('testing_type', ''))}</div>
                    </div>
                </div>
            </div>

            <!-- 3. Risk Rating Methodology -->
            <h2 class="section-heading" id="methodology"><span class="section-number">3</span> Risk Rating Methodology</h2>
            <div class="card">
                <p>Findings are rated <strong>Critical, High, Medium, Low,</strong> or <strong>Informational</strong> based on CVSS considerations, exploitability, and business impact. The overall risk score aggregates severity-weighted findings using the following formula:</p>
                <div class="finding-meta-grid" style="grid-template-columns: repeat(5, 1fr);">
                    <div style="text-align:center;">
                        <div style="font-size:1.3rem; font-weight:700; color:var(--sev-critical);">x25</div>
                        <div class="finding-field-label">Critical</div>
                    </div>
                    <div style="text-align:center;">
                        <div style="font-size:1.3rem; font-weight:700; color:var(--sev-high);">x15</div>
                        <div class="finding-field-label">High</div>
                    </div>
                    <div style="text-align:center;">
                        <div style="font-size:1.3rem; font-weight:700; color:var(--sev-medium);">x8</div>
                        <div class="finding-field-label">Medium</div>
                    </div>
                    <div style="text-align:center;">
                        <div style="font-size:1.3rem; font-weight:700; color:var(--sev-low);">x3</div>
                        <div class="finding-field-label">Low</div>
                    </div>
                    <div style="text-align:center;">
                        <div style="font-size:1.3rem; font-weight:700; color:var(--sev-info);">x1</div>
                        <div class="finding-field-label">Info</div>
                    </div>
                </div>
            </div>

            <!-- 4. Risk Matrix (CSS) -->
            <h2 class="section-heading" id="risk-matrix"><span class="section-number">4</span> Risk Matrix</h2>
            <div class="card">
                <p>The following matrix illustrates risk levels based on the combination of likelihood (exploitability) and impact:</p>
                <div class="risk-matrix">
                    <div class="risk-matrix-cell risk-matrix-corner">Impact / Likelihood</div>
                    <div class="risk-matrix-cell risk-matrix-header">Very Low</div>
                    <div class="risk-matrix-cell risk-matrix-header">Low</div>
                    <div class="risk-matrix-cell risk-matrix-header">Medium</div>
                    <div class="risk-matrix-cell risk-matrix-header">High</div>
                    <div class="risk-matrix-cell risk-matrix-header">Very High</div>

                    <div class="risk-matrix-cell risk-matrix-label">Critical</div>
                    <div class="risk-matrix-cell rm-medium">Medium</div>
                    <div class="risk-matrix-cell rm-high">High</div>
                    <div class="risk-matrix-cell rm-critical">Critical</div>
                    <div class="risk-matrix-cell rm-critical">Critical</div>
                    <div class="risk-matrix-cell rm-critical">Critical</div>

                    <div class="risk-matrix-cell risk-matrix-label">High</div>
                    <div class="risk-matrix-cell rm-low">Low</div>
                    <div class="risk-matrix-cell rm-medium">Medium</div>
                    <div class="risk-matrix-cell rm-high">High</div>
                    <div class="risk-matrix-cell rm-critical">Critical</div>
                    <div class="risk-matrix-cell rm-critical">Critical</div>

                    <div class="risk-matrix-cell risk-matrix-label">Medium</div>
                    <div class="risk-matrix-cell rm-low">Low</div>
                    <div class="risk-matrix-cell rm-low">Low</div>
                    <div class="risk-matrix-cell rm-medium">Medium</div>
                    <div class="risk-matrix-cell rm-high">High</div>
                    <div class="risk-matrix-cell rm-critical">Critical</div>

                    <div class="risk-matrix-cell risk-matrix-label">Low</div>
                    <div class="risk-matrix-cell rm-info">Info</div>
                    <div class="risk-matrix-cell rm-low">Low</div>
                    <div class="risk-matrix-cell rm-low">Low</div>
                    <div class="risk-matrix-cell rm-medium">Medium</div>
                    <div class="risk-matrix-cell rm-high">High</div>

                    <div class="risk-matrix-cell risk-matrix-label">Info</div>
                    <div class="risk-matrix-cell rm-info">Info</div>
                    <div class="risk-matrix-cell rm-info">Info</div>
                    <div class="risk-matrix-cell rm-low">Low</div>
                    <div class="risk-matrix-cell rm-low">Low</div>
                    <div class="risk-matrix-cell rm-medium">Medium</div>
                </div>
            </div>

            <!-- 5. Charts & Analytics -->
            <h2 class="section-heading" id="charts"><span class="section-number">5</span> Charts &amp; Analytics</h2>
            <div class="charts-grid">
                <div class="chart-card">
                    <div class="chart-card-title">Severity Distribution</div>
                    <div class="chart-container">
                        <canvas id="severityChart" width="300" height="250"></canvas>
                    </div>
                </div>
                <div class="chart-card">
                    <div class="chart-card-title">OWASP Category Breakdown</div>
                    <div class="chart-container">
                        <canvas id="owaspChart" width="350" height="250"></canvas>
                    </div>
                </div>
                <div class="chart-card">
                    <div class="chart-card-title">CWE Distribution</div>
                    <div class="chart-container">
                        <canvas id="cweChart" width="350" height="250"></canvas>
                    </div>
                </div>
                <div class="chart-card">
                    <div class="chart-card-title">Findings Trend Over Time</div>
                    <div class="chart-container">
                        <canvas id="trendChart" width="400" height="220"></canvas>
                    </div>
                </div>
            </div>
            <script>
            document.addEventListener('DOMContentLoaded', function() {{
                var sevData = {json.dumps([{"label": k.capitalize(), "value": v, "color": sev_colors[i % len(sev_colors)]} for i, (k, v) in enumerate(sorted(sev.items(), key=lambda x: _severity_order(x[0])))])};
                if (sevData.length) {{
                    new Chart(document.getElementById('severityChart'), {{
                        type: 'doughnut',
                        data: {{ labels: sevData.map(d => d.label), datasets: [{{ data: sevData.map(d => d.value), backgroundColor: sevData.map(d => d.color), borderWidth: 2, borderColor: '#fff' }}] }},
                        options: {{ responsive: true, plugins: {{ legend: {{ position: 'bottom', labels: {{ padding: 16, usePointStyle: true, pointStyle: 'circle' }} }} }} }}
                    }});
                }}
                var owaspData = {json.dumps([{"label": k[:30] if len(k) > 30 else k, "value": v} for k, v in list(owasp.items())[:8]])};
                if (owaspData.length) {{
                    new Chart(document.getElementById('owaspChart'), {{
                        type: 'bar',
                        data: {{ labels: owaspData.map(d => d.label), datasets: [{{ label: 'Count', data: owaspData.map(d => d.value), backgroundColor: 'rgba(30,58,95,0.85)', borderRadius: 6, borderSkipped: false }}] }},
                        options: {{ indexAxis: 'y', responsive: true, plugins: {{ legend: {{ display: false }} }}, scales: {{ x: {{ grid: {{ display: false }}, ticks: {{ stepSize: 1 }} }}, y: {{ grid: {{ display: false }} }} }} }}
                    }});
                }}
                var cweData = {json.dumps([{"label": k[:20] if len(str(k)) > 20 else k, "value": v} for k, v in list(cwe.items())[:10] if str(k) != "N/A"])};
                if (cweData.length) {{
                    new Chart(document.getElementById('cweChart'), {{
                        type: 'bar',
                        data: {{ labels: cweData.map(d => d.label), datasets: [{{ label: 'Count', data: cweData.map(d => d.value), backgroundColor: 'rgba(22,163,74,0.75)', borderRadius: 6 }}] }},
                        options: {{ indexAxis: 'y', responsive: true, plugins: {{ legend: {{ display: false }} }}, scales: {{ x: {{ ticks: {{ stepSize: 1 }} }}, y: {{ grid: {{ display: false }} }} }} }}
                    }});
                }}
                var trendLabels = {json.dumps(trend_dates)};
                var trendValues = {json.dumps(trend_counts)};
                if (trendLabels.length && document.getElementById('trendChart')) {{
                    new Chart(document.getElementById('trendChart'), {{
                        type: 'line',
                        data: {{ labels: trendLabels, datasets: [{{ label: 'Findings', data: trendValues, borderColor: '#1e3a5f', backgroundColor: 'rgba(37,99,235,0.1)', fill: true, tension: 0.3, pointRadius: 4 }}] }},
                        options: {{ responsive: true, plugins: {{ legend: {{ display: false }} }}, scales: {{ y: {{ beginAtZero: true, ticks: {{ stepSize: 1 }} }} }} }}
                    }});
                }}
            }});
            </script>

            <!-- 6. Summary Statistics -->
            <h2 class="section-heading" id="stats"><span class="section-number">6</span> Summary Statistics</h2>
            <table><tr><th>Severity</th><th>Count</th></tr>{sev_rows}</table>

            <!-- 7. OWASP Top 10 Mapping -->
            <h2 class="section-heading" id="owasp"><span class="section-number">7</span> OWASP Top 10 Mapping</h2>
            <table><tr><th>Category</th><th>Count</th></tr>{owasp_rows}</table>

            <!-- 8. CWE Mapping -->
            <h2 class="section-heading" id="cwe"><span class="section-number">8</span> CWE Mapping</h2>
            <table><tr><th>CWE ID</th><th>Count</th></tr>{cwe_rows}</table>
            <h3>CWE to Compliance Framework Cross-Reference</h3>
            <table>
                <tr><th>CWE</th><th>Finding</th><th>Severity</th><th>MITRE ATT&amp;CK</th><th>ISO 27001</th><th>NIST 800-53</th></tr>
                {cwe_detail_rows}
            </table>

            <!-- 9. Compliance Frameworks -->
            <h2 class="section-heading" id="compliance"><span class="section-number">9</span> Compliance Frameworks</h2>
            <div class="card">
                <p>All findings are mapped to the following industry-standard compliance frameworks:</p>
                <table>
                    <tr><th>Framework</th><th>Description</th><th>Status</th></tr>
                    {compliance_rows}
                </table>
            </div>

            <!-- 10. Findings Summary -->
            <h2 class="section-heading" id="findings-table"><span class="section-number">10</span> Findings Summary</h2>
            <table>
                <tr><th>#</th><th>Title</th><th>Severity</th><th>OWASP</th><th>CWE</th><th>Affected URL</th></tr>
                {find_rows}
            </table>

            <!-- 11. Finding Details with Evidence -->
            <h2 class="section-heading" id="finding-details"><span class="section-number">11</span> Finding Details with Evidence</h2>
"""
    for i, f in enumerate(findings, 1):
        comp = f.get("compliance") or {}
        comp_tags = []
        if f.get("owasp_category"):
            comp_tags.append(f"OWASP: {f['owasp_category']}")
        if f.get("cwe_id"):
            comp_tags.append(f"{f['cwe_id']}")
        if comp.get("mitre_attack"):
            comp_tags.append(f"ATT&CK: {comp['mitre_attack']}")
        if comp.get("iso_27001"):
            comp_tags.append(f"ISO 27001: {comp['iso_27001']}")
        if comp.get("nist_800_53"):
            comp_tags.append(f"NIST: {comp['nist_800_53']}")
        comp_tags_html = "".join(f'<span class="compliance-tag">{_html_escape(t)}</span>' for t in comp_tags)

        evidence_html = ""
        evidence_list = f.get("evidence") or []
        for ev_idx, ev in enumerate(evidence_list):
            if isinstance(ev, dict):
                b64 = _evidence_to_base64(project_id, ev)
                if b64:
                    fn = ev.get("filename", "Evidence")
                    evidence_html += f"""
                    <div class="evidence-container">
                        <div class="evidence-caption">Evidence {ev_idx + 1}: {_html_escape(fn)}</div>
                        <div class="evidence-img-wrap">
                            <img src="{b64}" alt="{_html_escape(fn)}" class="evidence-img" />
                        </div>
                    </div>"""

        sev_lower = f.get("severity", "").lower()
        cvss_val = f.get("cvss_score") or "N/A"

        req_resp_html = ""
        if f.get("request"):
            req_resp_html += f'<div class="finding-field"><div class="finding-field-label">Request (Automated Evidence)</div><pre class="evidence-pre">{_html_escape(str(f.get("request", "")))}</pre></div>'
        if f.get("response"):
            req_resp_html += f'<div class="finding-field"><div class="finding-field-label">Response (Automated Evidence)</div><pre class="evidence-pre">{_html_escape(str(f.get("response", "")))}</pre></div>'

        html += f"""
            <div class="finding-block" id="finding-{i}">
                <div class="finding-block-header sev-bg-{sev_lower}">
                    <h3 class="finding-block-title">Finding #{i}: {_html_escape(f.get('title',''))}</h3>
                    {_sev_badge(f.get('severity',''))}
                </div>
                <div class="finding-block-body">
                    <div class="finding-meta-grid">
                        <div>
                            <div class="finding-field-label">Affected URL</div>
                            <div class="finding-field-value"><code>{_html_escape(f.get('affected_url') or '-')}</code></div>
                        </div>
                        <div>
                            <div class="finding-field-label">CVSS Score</div>
                            <div class="finding-field-value"><strong>{_html_escape(str(cvss_val))}</strong></div>
                        </div>
                        <div>
                            <div class="finding-field-label">OWASP Category</div>
                            <div class="finding-field-value">{_html_escape(f.get('owasp_category') or '-')}</div>
                        </div>
                        <div>
                            <div class="finding-field-label">CWE ID</div>
                            <div class="finding-field-value"><code>{_html_escape(f.get('cwe_id') or '-')}</code></div>
                        </div>
                    </div>
                    <div class="finding-field">
                        <div class="finding-field-label">Description</div>
                        <div class="finding-field-value">{_html_escape(f.get('description') or '-').replace(chr(10), '<br>')}</div>
                    </div>
                    <div class="finding-field">
                        <div class="finding-field-label">Impact</div>
                        <div class="finding-field-value">{_html_escape(f.get('impact') or '-').replace(chr(10), '<br>')}</div>
                    </div>
                    <div class="finding-field">
                        <div class="finding-field-label">Reproduction Steps</div>
                        <div class="finding-field-value">{_html_escape(f.get('reproduction_steps') or '-').replace(chr(10), '<br>')}</div>
                    </div>
                    {evidence_html}
                    {req_resp_html}
                    <div class="finding-field">
                        <div class="finding-field-label">Recommendation</div>
                        <div class="finding-field-value">{_html_escape(f.get('recommendation') or '-').replace(chr(10), '<br>')}</div>
                    </div>
                    {f'<div class="finding-field"><div class="finding-field-label">Compliance Mapping</div><div class="compliance-tags">{comp_tags_html}</div></div>' if comp_tags_html else ''}
                </div>
            </div>
"""

    # Timeline section
    html += f"""
            <!-- 12. Remediation Timeline -->
            <h2 class="section-heading" id="timeline"><span class="section-number">12</span> Remediation Timeline</h2>
            <div class="card">
                <p>The following timeline outlines recommended remediation targets based on finding severity:</p>
                <div class="finding-meta-grid" style="grid-template-columns: repeat(4, 1fr); margin-bottom: 1.5rem;">
                    <div style="text-align:center; padding:0.75rem; background:var(--sev-critical-bg); border-radius:var(--radius-xs);">
                        <div style="font-weight:700; color:var(--sev-critical);">0-7 days</div>
                        <div class="finding-field-label">Critical</div>
                    </div>
                    <div style="text-align:center; padding:0.75rem; background:var(--sev-high-bg); border-radius:var(--radius-xs);">
                        <div style="font-weight:700; color:var(--sev-high);">7-30 days</div>
                        <div class="finding-field-label">High</div>
                    </div>
                    <div style="text-align:center; padding:0.75rem; background:var(--sev-medium-bg); border-radius:var(--radius-xs);">
                        <div style="font-weight:700; color:var(--sev-medium);">30-90 days</div>
                        <div class="finding-field-label">Medium</div>
                    </div>
                    <div style="text-align:center; padding:0.75rem; background:var(--sev-low-bg); border-radius:var(--radius-xs);">
                        <div style="font-weight:700; color:var(--sev-low);">Next Release</div>
                        <div class="finding-field-label">Low / Info</div>
                    </div>
                </div>
                <div class="timeline">
                    {timeline_html}
                </div>
            </div>

        </div><!-- end report-body -->

        <!-- Footer -->
        <div class="report-footer">
            <div class="footer-confidentiality">
                <strong>Confidentiality Notice</strong>
                This document contains confidential and privileged information intended solely for the authorized recipient(s).
                Any unauthorized review, use, disclosure, or distribution is strictly prohibited. If you have received this
                document in error, please notify the sender immediately and destroy all copies. This report and its contents
                are protected under applicable confidentiality agreements and information security policies.
            </div>
            <div class="footer-meta">
                <span>Report generated by {_html_escape(org_name)}</span>
                <span>{gen}</span>
                <span>Classification: {_html_escape(p.get('classification') or 'Confidential')}</span>
            </div>
        </div>
    </div><!-- end report-container -->
"""
    footer_prefix = f'<!-- Integrity: SHA-256: '
    content_to_hash = html + footer_prefix
    report_hash = hashlib.sha256(content_to_hash.encode("utf-8")).hexdigest()
    html += footer_prefix + report_hash + " -->\n</body>\n</html>\n"
    return html


def generate_docx(data: dict) -> bytes:
    """Generate Big 4 style professional DOCX report with embedded screenshots."""
    org = data.get("organization") or {"name": "AppSecD", "logo_base64": None}
    org_name = org.get("name") or "AppSecD"
    from docx.shared import RGBColor, Cm
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    doc = Document()
    p = data["project"]
    findings = data["findings"]
    risk = data["risk_score"]
    risk_level = data.get("risk_level", "Medium")
    cov = data["coverage_pct"]
    sev = data["severity_distribution"]
    owasp = data["owasp_mapping"]
    cwe = data["cwe_mapping"]
    project_id = data.get("project_id", "")
    gen = data["generated_at"]

    sev_colors = {
        "critical": "DC2626",
        "high": "EA580C",
        "medium": "D97706",
        "low": "16A34A",
        "info": "2563EB",
    }
    primary_color = "1E3A5F"
    primary_light = "2563EB"

    # --- Style helpers ---
    def _set_cell_shading(cell, color_hex: str):
        """Set background shading for a table cell."""
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    def _set_cell_text(cell, text: str, bold: bool = False, size: int = 9,
                       color: str = None, alignment=None):
        """Set cell text with formatting."""
        cell.text = ""
        para = cell.paragraphs[0]
        if alignment:
            para.alignment = alignment
        run = para.add_run(str(text))
        run.bold = bold
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
        run.font.name = "Calibri"

    def _add_styled_table(doc, headers: list, rows_data: list, col_widths: list = None):
        """Create a professionally styled table."""
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        hdr = table.rows[0]
        for i, header_text in enumerate(headers):
            cell = hdr.cells[i]
            _set_cell_shading(cell, primary_color)
            _set_cell_text(cell, header_text, bold=True, size=9, color="FFFFFF")

        # Data rows
        for row_idx, row_data in enumerate(rows_data):
            row = table.add_row()
            bg = "F7FAFC" if row_idx % 2 == 0 else "FFFFFF"
            for col_idx, cell_text in enumerate(row_data):
                cell = row.cells[col_idx]
                _set_cell_shading(cell, bg)
                is_bold = col_idx == 0
                _set_cell_text(cell, str(cell_text), bold=is_bold, size=9)
        return table

    def _add_heading_styled(doc, text: str, level: int = 1):
        """Add a heading with custom styling."""
        heading = doc.add_heading(text, level=level)
        for run in heading.runs:
            run.font.color.rgb = RGBColor.from_string(primary_color)
            run.font.name = "Calibri"
        return heading

    def _add_field_para(doc, label: str, value: str):
        """Add a label: value paragraph."""
        para = doc.add_paragraph()
        run_label = para.add_run(f"{label}: ")
        run_label.bold = True
        run_label.font.size = Pt(10)
        run_label.font.color.rgb = RGBColor.from_string("4A5568")
        run_label.font.name = "Calibri"
        run_value = para.add_run(str(value))
        run_value.font.size = Pt(10)
        run_value.font.name = "Calibri"
        return para

    # --- Modify default styles ---
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)
    font.color.rgb = RGBColor.from_string("1A202C")

    for level in range(0, 4):
        style_name = f"Heading {level}" if level > 0 else "Title"
        try:
            hs = doc.styles[style_name]
            hs.font.name = "Calibri"
            hs.font.color.rgb = RGBColor.from_string(primary_color)
        except Exception:
            pass

    # ============================================================
    # Cover Page
    # ============================================================
    doc.add_paragraph()  # Spacer
    doc.add_paragraph()

    # Brand name
    brand = doc.add_paragraph()
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = brand.add_run(org_name.upper())
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(primary_light)
    run.font.name = "Calibri"
    run.bold = False

    doc.add_paragraph()

    # Title
    title = doc.add_heading("Application Security Assessment Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(26)
        run.font.color.rgb = RGBColor.from_string(primary_color)
        run.font.name = "Calibri"

    # Subtitle (application name)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(p.get("application_name", "") or p.get("name", ""))
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor.from_string(primary_light)
    run.font.name = "Calibri"
    run.bold = True

    doc.add_paragraph()

    # Cover metadata table
    cover_table = doc.add_table(rows=6, cols=2)
    cover_table.style = "Table Grid"
    cover_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cover_data = [
        ("Target URL", p.get("application_url", "")),
        ("Environment", p.get("environment", "")),
        ("Testing Type", p.get("testing_type", "")),
        ("Classification", p.get("classification") or "Confidential"),
        ("Risk Level", f"{risk_level} ({risk}/100)"),
        ("Report Generated", gen),
    ]
    for i, (label, value) in enumerate(cover_data):
        _set_cell_shading(cover_table.rows[i].cells[0], "F0F4F8")
        _set_cell_text(cover_table.rows[i].cells[0], label, bold=True, size=10, color="4A5568")
        _set_cell_text(cover_table.rows[i].cells[1], str(value), size=10)

    doc.add_page_break()

    # ============================================================
    # Table of Contents placeholder
    # ============================================================
    toc_heading = _add_heading_styled(doc, "Table of Contents", level=1)
    toc_items = [
        "1. Executive Summary",
        "2. Scope & Methodology",
        "3. Risk Score Assessment",
        "4. Severity Distribution",
        "5. OWASP Top 10 Mapping",
        "6. CWE Mapping Table",
        "7. Compliance Framework Mapping",
        "8. Detailed Findings with Evidence",
        "9. Remediation Timeline",
        "10. Confidentiality Notice",
    ]
    for item in toc_items:
        toc_para = doc.add_paragraph()
        run = toc_para.add_run(item)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor.from_string(primary_light)
        run.font.name = "Calibri"
    doc.add_page_break()

    # ============================================================
    # 1. Executive Summary
    # ============================================================
    _add_heading_styled(doc, "1. Executive Summary", level=1)
    ai = data.get("ai_report_content") or {}
    if ai:
        exec_sum = ai.get("executive_summary") or ai.get("ai_summary")
        if exec_sum:
            p_ai = doc.add_paragraph()
            run = p_ai.add_run("AI-Generated Executive Summary")
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor.from_string(primary_light)
            doc.add_paragraph(str(exec_sum), style="Normal")
            doc.add_paragraph()
        tech = ai.get("technical_summary")
        if tech:
            p_tech = doc.add_paragraph()
            run = p_tech.add_run("Technical Summary")
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor.from_string(primary_light)
            doc.add_paragraph(str(tech), style="Normal")
            doc.add_paragraph()
        strat = ai.get("strategic_recommendations")
        if strat:
            p_strat = doc.add_paragraph()
            run = p_strat.add_run("Strategic Recommendations")
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor.from_string("0ea5e9")
            strat_text = "\n".join(f"• {s}" for s in strat) if isinstance(strat, list) else str(strat)
            doc.add_paragraph(strat_text, style="Normal")
            doc.add_paragraph()
    doc.add_paragraph(
        f"This report presents the findings of a security assessment conducted on "
        f"{p.get('application_name', '')}. The assessment evaluated the application against "
        f"industry-standard security controls including OWASP Top 10, CWE Top 25, "
        f"MITRE ATT&CK, ISO 27001, and NIST 800-53 frameworks."
    )
    doc.add_paragraph()

    # Risk score KPI table
    kpi_table = doc.add_table(rows=2, cols=4)
    kpi_table.style = "Table Grid"
    kpi_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    kpi_headers = ["Risk Score", "Total Findings", "Test Coverage", "Critical/High"]
    critical_high = sev.get("critical", 0) + sev.get("high", 0)
    kpi_values = [f"{risk}/100", str(len(findings)), f"{cov}%", str(critical_high)]
    kpi_colors_list = [
        sev_colors.get(risk_level.lower(), "1E3A5F"),
        primary_light,
        "0EA5E9",
        sev_colors["critical"],
    ]
    for i in range(4):
        _set_cell_shading(kpi_table.rows[0].cells[i], primary_color)
        _set_cell_text(kpi_table.rows[0].cells[i], kpi_headers[i], bold=True, size=9, color="FFFFFF",
                       alignment=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_text(kpi_table.rows[1].cells[i], kpi_values[i], bold=True, size=14,
                       color=kpi_colors_list[i], alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    # Risk level statement
    risk_para = doc.add_paragraph()
    run = risk_para.add_run(f"Overall Risk Rating: {risk_level} ")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Calibri"
    run = risk_para.add_run(f"({risk}/100)")
    run.font.size = Pt(12)
    run.font.name = "Calibri"

    doc.add_paragraph(
        f"Test Coverage: {cov}% ({p.get('tested_count', 0)}/{p.get('total_test_cases', 0)} test cases executed)."
    )

    if findings:
        critical_high_count = sum(1 for x in findings if x.get("severity", "").lower() in ("critical", "high"))
        if critical_high_count:
            warn_para = doc.add_paragraph()
            run = warn_para.add_run(f"{critical_high_count} Critical/High severity finding(s) require immediate remediation.")
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(sev_colors["critical"])
            run.font.size = Pt(10)
            run.font.name = "Calibri"

    doc.add_paragraph()
    rec_para = doc.add_paragraph()
    run = rec_para.add_run("Recommendation: ")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    run = rec_para.add_run(
        "Address Critical and High severity findings within 30 days. "
        "Medium and Low findings should be remediated in subsequent release cycles."
    )
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    doc.add_paragraph()

    # ============================================================
    # 2. Scope & Methodology
    # ============================================================
    _add_heading_styled(doc, "2. Scope & Methodology", level=1)
    scope_text = p.get("testing_scope") or "In-scope: Application under test. Out-of-scope: Third-party systems."
    meth_type = "Black-box" if p.get("testing_type") == "black_box" else "Grey-box" if p.get("testing_type") == "grey_box" else "White-box"

    scope_table = doc.add_table(rows=4, cols=2)
    scope_table.style = "Table Grid"
    scope_data = [
        ("Scope", scope_text),
        ("Methodology", f"{meth_type} testing in {p.get('environment', 'staging')} environment. Manual and semi-automated techniques per OWASP Testing Guide."),
        ("Target Completion", p.get("target_completion_date") or "N/A"),
        ("Testing Standard", "OWASP Testing Guide v4.2, OWASP WSTG"),
    ]
    for i, (label, value) in enumerate(scope_data):
        bg = "F0F4F8" if i % 2 == 0 else "FFFFFF"
        _set_cell_shading(scope_table.rows[i].cells[0], bg)
        _set_cell_shading(scope_table.rows[i].cells[1], bg)
        _set_cell_text(scope_table.rows[i].cells[0], label, bold=True, size=10, color="4A5568")
        _set_cell_text(scope_table.rows[i].cells[1], str(value), size=10)
    doc.add_paragraph()

    # ============================================================
    # 3. Risk Score Assessment
    # ============================================================
    _add_heading_styled(doc, "3. Risk Score Assessment", level=1)
    doc.add_paragraph(
        "The risk score is calculated using a weighted severity model. "
        "Each finding contributes to the overall score based on its severity level:"
    )

    weight_table = doc.add_table(rows=1, cols=5)
    weight_table.style = "Table Grid"
    weight_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    weights = [("Critical", "x25", "critical"), ("High", "x15", "high"),
               ("Medium", "x8", "medium"), ("Low", "x3", "low"), ("Info", "x1", "info")]
    for i, (label, multiplier, sev_key) in enumerate(weights):
        cell = weight_table.rows[0].cells[i]
        _set_cell_shading(cell, sev_colors[sev_key])
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(f"{label}\n{multiplier}")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string("FFFFFF")
        run.font.name = "Calibri"

    doc.add_paragraph()
    doc.add_paragraph(
        f"Based on this methodology, the current risk score is {risk}/100, "
        f"which corresponds to a {risk_level} risk level."
    )

    # Visual risk scale table
    scale_table = doc.add_table(rows=2, cols=4)
    scale_table.style = "Table Grid"
    scale_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    scale_items = [
        ("Low (0-24)", sev_colors["low"], risk_level == "Low"),
        ("Medium (25-49)", sev_colors["medium"], risk_level == "Medium"),
        ("High (50-74)", sev_colors["high"], risk_level == "High"),
        ("Critical (75-100)", sev_colors["critical"], risk_level == "Critical"),
    ]
    for i, (label, color, is_active) in enumerate(scale_items):
        # Top row: color indicator
        _set_cell_shading(scale_table.rows[0].cells[i], color if is_active else "E2E8F0")
        _set_cell_text(scale_table.rows[0].cells[i], ">>>>" if is_active else "",
                       bold=True, size=10, color="FFFFFF" if is_active else "A0AEC0",
                       alignment=WD_ALIGN_PARAGRAPH.CENTER)
        # Bottom row: label
        _set_cell_shading(scale_table.rows[1].cells[i], "F7FAFC")
        _set_cell_text(scale_table.rows[1].cells[i], label, bold=is_active, size=9,
                       color=color if is_active else "718096",
                       alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    # ============================================================
    # 4. Severity Distribution
    # ============================================================
    _add_heading_styled(doc, "4. Severity Distribution", level=1)
    sev_data = [(k.capitalize(), str(v)) for k, v in sorted(sev.items(), key=lambda x: _severity_order(x[0]))]
    _add_styled_table(doc, ["Severity", "Count"], sev_data)
    doc.add_paragraph()

    # ============================================================
    # 5. OWASP Top 10 Mapping
    # ============================================================
    _add_heading_styled(doc, "5. OWASP Top 10 Mapping", level=1)
    owasp_data = [(str(k), str(v)) for k, v in sorted(owasp.items())]
    _add_styled_table(doc, ["OWASP Category", "Count"], owasp_data)
    doc.add_paragraph()

    # ============================================================
    # 6. CWE Mapping Table
    # ============================================================
    _add_heading_styled(doc, "6. CWE Mapping Table", level=1)
    doc.add_paragraph(
        "The following table maps each finding to its CWE identifier and "
        "cross-references the corresponding compliance frameworks."
    )
    cwe_table_data = []
    for f in findings:
        comp = f.get("compliance") or {}
        cwe_table_data.append((
            f.get("cwe_id") or "-",
            f.get("title", ""),
            f.get("severity", ""),
            comp.get("mitre_attack") or "-",
            comp.get("iso_27001") or "-",
            comp.get("nist_800_53") or "-",
        ))
    cwe_full_table = _add_styled_table(
        doc,
        ["CWE ID", "Finding", "Severity", "MITRE ATT&CK", "ISO 27001", "NIST 800-53"],
        cwe_table_data
    )

    # Color-code severity cells in the CWE table
    for row_idx, row_data in enumerate(cwe_table_data):
        sev_val = row_data[2].lower()
        if sev_val in sev_colors:
            sev_cell = cwe_full_table.rows[row_idx + 1].cells[2]
            _set_cell_text(sev_cell, row_data[2], bold=True, size=9, color=sev_colors[sev_val])

    doc.add_paragraph()

    # Summary CWE count table
    cwe_count_data = [(str(k), str(v)) for k, v in sorted(cwe.items())]
    if cwe_count_data:
        sub_h = doc.add_heading("CWE Summary Counts", level=2)
        for run in sub_h.runs:
            run.font.color.rgb = RGBColor.from_string(primary_color)
            run.font.name = "Calibri"
        _add_styled_table(doc, ["CWE ID", "Count"], cwe_count_data)
    doc.add_paragraph()

    # ============================================================
    # 7. Compliance Framework Mapping
    # ============================================================
    _add_heading_styled(doc, "7. Compliance Framework Mapping", level=1)
    doc.add_paragraph(
        "All findings have been mapped to the following industry-standard compliance "
        "frameworks to support governance, risk, and compliance (GRC) requirements."
    )

    comp_data = [
        ("OWASP Top 10 (2021)", "Web Application Security Standard", "Mapped"),
        ("CWE Top 25", "Common Weakness Enumeration", "Mapped"),
        ("MITRE ATT&CK", "Adversarial Tactics and Techniques", "Mapped"),
        ("ISO 27001 Annex A", "Information Security Controls", "Mapped"),
        ("NIST 800-53", "Security and Privacy Controls", "Mapped"),
    ]
    comp_table = _add_styled_table(doc, ["Framework", "Description", "Status"], comp_data)

    # Color the status cells green
    for row_idx in range(len(comp_data)):
        status_cell = comp_table.rows[row_idx + 1].cells[2]
        _set_cell_shading(status_cell, "DCFCE7")
        _set_cell_text(status_cell, "Mapped", bold=True, size=9, color="166534",
                       alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    # ============================================================
    # 8. Detailed Findings with Evidence
    # ============================================================
    _add_heading_styled(doc, "8. Detailed Findings with Evidence", level=1)

    for i, f in enumerate(findings, 1):
        sev_val = f.get("severity", "").lower()
        sev_color = sev_colors.get(sev_val, "718096")

        # Finding heading with severity color
        finding_heading = doc.add_heading(level=2)
        run = finding_heading.add_run(f"Finding #{i}: {f.get('title', '')}")
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor.from_string(primary_color)
        run = finding_heading.add_run(f"  [{f.get('severity', '')}]")
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor.from_string(sev_color)
        run.bold = True

        # Finding metadata table
        meta_table = doc.add_table(rows=4, cols=2)
        meta_table.style = "Table Grid"
        meta_items = [
            ("Affected URL", f.get("affected_url") or "-"),
            ("CVSS Score", str(f.get("cvss_score") or "N/A")),
            ("OWASP Category", f.get("owasp_category") or "-"),
            ("CWE ID", f.get("cwe_id") or "-"),
        ]
        for mi, (label, value) in enumerate(meta_items):
            bg = "F0F4F8" if mi % 2 == 0 else "FFFFFF"
            _set_cell_shading(meta_table.rows[mi].cells[0], bg)
            _set_cell_shading(meta_table.rows[mi].cells[1], bg)
            _set_cell_text(meta_table.rows[mi].cells[0], label, bold=True, size=9, color="4A5568")
            _set_cell_text(meta_table.rows[mi].cells[1], str(value), size=9)

        doc.add_paragraph()

        # Description
        _add_field_para(doc, "Description", f.get("description") or "-")

        # Impact
        _add_field_para(doc, "Impact", f.get("impact") or "-")

        # Reproduction Steps
        _add_field_para(doc, "Reproduction Steps", f.get("reproduction_steps") or "-")

        # Evidence
        evidence_list = f.get("evidence") or []
        for ev_idx, ev in enumerate(evidence_list):
            if isinstance(ev, dict):
                fpath = _get_evidence_file_path(project_id, ev)
                if fpath and fpath.exists() and fpath.suffix.lower() in {".png", ".jpg", ".jpeg", ".gif", ".webp"}:
                    try:
                        fn = ev.get("filename", "Screenshot")
                        ev_para = doc.add_paragraph()
                        run = ev_para.add_run(f"Evidence {ev_idx + 1}: {fn}")
                        run.bold = True
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor.from_string(primary_color)
                        run.font.name = "Calibri"
                        doc.add_picture(str(fpath), width=Inches(5.5))
                        doc.add_paragraph()
                    except Exception:
                        pass

        # Request/Response (Automated Evidence)
        if f.get("request"):
            _add_field_para(doc, "Request (Automated Evidence)", f.get("request") or "-")
        if f.get("response"):
            _add_field_para(doc, "Response (Automated Evidence)", f.get("response") or "-")

        # Recommendation
        _add_field_para(doc, "Recommendation", f.get("recommendation") or "-")

        # Compliance mapping tags
        comp = f.get("compliance") or {}
        comp_parts = []
        if f.get("owasp_category"):
            comp_parts.append(f"OWASP: {f['owasp_category']}")
        if f.get("cwe_id"):
            comp_parts.append(f"{f['cwe_id']}")
        if comp.get("mitre_attack"):
            comp_parts.append(f"ATT&CK: {comp['mitre_attack']}")
        if comp.get("iso_27001"):
            comp_parts.append(f"ISO 27001: {comp['iso_27001']}")
        if comp.get("nist_800_53"):
            comp_parts.append(f"NIST: {comp['nist_800_53']}")
        if comp_parts:
            comp_para = doc.add_paragraph()
            run = comp_para.add_run("Compliance Mapping: ")
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string("4A5568")
            run.font.name = "Calibri"
            run = comp_para.add_run(" | ".join(comp_parts))
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string(primary_light)
            run.font.name = "Calibri"

        doc.add_paragraph()  # Spacer between findings

    # ============================================================
    # 9. Remediation Timeline
    # ============================================================
    _add_heading_styled(doc, "9. Remediation Timeline", level=1)
    doc.add_paragraph(
        "The following table outlines recommended remediation timeframes based on finding severity:"
    )

    timeline_data = []
    for i, f in enumerate(findings, 1):
        sev_val = f.get("severity", "").lower()
        if sev_val == "critical":
            target = "Immediate (0-7 days)"
        elif sev_val == "high":
            target = "Urgent (7-30 days)"
        elif sev_val == "medium":
            target = "Planned (30-90 days)"
        else:
            target = "Backlog (Next release)"
        timeline_data.append((
            str(i),
            f.get("title", ""),
            f.get("severity", ""),
            target,
        ))

    timeline_table = _add_styled_table(
        doc,
        ["#", "Finding", "Severity", "Remediation Target"],
        timeline_data
    )

    # Color-code severity cells in timeline
    for row_idx, row_data in enumerate(timeline_data):
        sev_val = row_data[2].lower()
        if sev_val in sev_colors:
            sev_cell = timeline_table.rows[row_idx + 1].cells[2]
            _set_cell_text(sev_cell, row_data[2], bold=True, size=9, color=sev_colors[sev_val])
    doc.add_paragraph()

    # ============================================================
    # 10. Confidentiality Notice
    # ============================================================
    _add_heading_styled(doc, "10. Confidentiality Notice", level=1)

    notice_para = doc.add_paragraph()
    run = notice_para.add_run("CONFIDENTIALITY NOTICE")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(primary_color)
    run.font.name = "Calibri"

    doc.add_paragraph(
        "This document contains confidential and privileged information intended solely for the "
        "authorized recipient(s). Any unauthorized review, use, disclosure, or distribution is "
        "strictly prohibited. If you have received this document in error, please notify the sender "
        "immediately and destroy all copies. This report and its contents are protected under "
        "applicable confidentiality agreements and information security policies."
    )

    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    run = footer_para.add_run(f"Report generated by {org_name} at {gen}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string("718096")
    run.font.name = "Calibri"

    classification_para = doc.add_paragraph()
    classification_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = classification_para.add_run(f"--- {(p.get('classification') or 'Confidential').upper()} ---")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(primary_color)
    run.font.name = "Calibri"

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def generate_pdf(data: dict) -> bytes:
    """Generate Big 4 style professional PDF report with embedded screenshots."""
    if not FPDF:
        raise RuntimeError("fpdf2 not installed")
    org = data.get("organization") or {"name": "AppSecD"}
    org_name = org.get("name") or "AppSecD"
    p = data["project"]
    findings = data["findings"]
    risk = data["risk_score"]
    risk_level = data.get("risk_level", "Medium")
    cov = data["coverage_pct"]
    sev = data["severity_distribution"]
    owasp = data["owasp_mapping"]
    cwe = data["cwe_mapping"]
    project_id = data.get("project_id", "")
    gen = data["generated_at"]

    sev_color_map = {
        "critical": (220, 38, 38),
        "high": (234, 88, 12),
        "medium": (217, 119, 6),
        "low": (22, 163, 106),
        "info": (37, 99, 235),
    }
    primary_color = (30, 58, 95)
    primary_light = (37, 99, 235)
    accent_color = (14, 165, 233)

    def safe_text(s, max_len: int = 500) -> str:
        t = str(s or "-").replace("\n", " ").replace("\r", "")[:max_len]
        return t if t else "-"

    class ReportPDF(FPDF):
        def header(self):
            if self.page_no() == 1:
                return  # Cover page has no header
            # Top accent line
            self.set_fill_color(*primary_color)
            self.rect(0, 0, 210, 3, "F")
            self.set_fill_color(*primary_light)
            self.rect(0, 3, 210, 1, "F")
            self.ln(6)
            self.set_font("Helvetica", "B", 9)
            self.set_text_color(*primary_color)
            self.cell(95, 5, f"{org_name} Security Assessment Report", ln=0, align="L")
            self.set_font("Helvetica", "", 8)
            self.set_text_color(100, 100, 100)
            self.cell(95, 5, f"{p.get('application_name', '')} | {p.get('classification') or 'Confidential'}", ln=True, align="R")
            self.set_draw_color(226, 232, 240)
            self.line(10, self.get_y() + 1, 200, self.get_y() + 1)
            self.ln(4)
            self.set_text_color(0, 0, 0)

        def footer(self):
            self.set_y(-18)
            self.set_draw_color(226, 232, 240)
            self.line(10, self.get_y(), 200, self.get_y())
            self.ln(2)
            self.set_font("Helvetica", "", 7)
            self.set_text_color(120, 120, 120)
            self.cell(95, 5, f"Generated: {gen[:10]} | {org_name}", ln=0, align="L")
            self.cell(95, 5, f"Page {self.page_no()}/{{nb}} | Confidential", ln=True, align="R")
            self.set_text_color(0, 0, 0)

        def section_heading(self, num: int, title: str):
            self.ln(4)
            # Section number circle
            self.set_fill_color(*primary_light)
            x_start = self.get_x()
            y_start = self.get_y()
            self.set_text_color(255, 255, 255)
            self.set_font("Helvetica", "B", 9)
            self.cell(8, 8, str(num), align="C", fill=True)
            self.set_text_color(*primary_color)
            self.set_font("Helvetica", "B", 13)
            self.cell(0, 8, f"  {title}", ln=True)
            self.set_draw_color(*primary_light)
            self.set_line_width(0.5)
            self.line(10, self.get_y() + 1, 200, self.get_y() + 1)
            self.set_line_width(0.2)
            self.ln(4)
            self.set_text_color(0, 0, 0)

        def severity_badge(self, severity: str):
            s = severity.lower()
            color = sev_color_map.get(s, (100, 100, 100))
            self.set_fill_color(*color)
            self.set_text_color(255, 255, 255)
            self.set_font("Helvetica", "B", 8)
            w = self.get_string_width(severity.upper()) + 6
            self.cell(w, 5, severity.upper(), align="C", fill=True)
            self.set_text_color(0, 0, 0)

        def risk_gauge(self, score: float, y_offset: float = 0):
            """Draw a visual risk score gauge bar."""
            x = 10
            y = self.get_y() + y_offset
            bar_w = 120
            bar_h = 8
            # Background
            self.set_fill_color(226, 232, 240)
            self.rect(x, y, bar_w, bar_h, "F")
            # Fill
            fill_w = (score / 100) * bar_w
            if score >= 75:
                self.set_fill_color(*sev_color_map["critical"])
            elif score >= 50:
                self.set_fill_color(*sev_color_map["high"])
            elif score >= 25:
                self.set_fill_color(*sev_color_map["medium"])
            else:
                self.set_fill_color(*sev_color_map["low"])
            self.rect(x, y, fill_w, bar_h, "F")
            # Label
            self.set_font("Helvetica", "B", 9)
            self.set_xy(x + bar_w + 4, y)
            self.cell(30, bar_h, f"{score}/100", align="L")

        def key_value_row(self, label: str, value: str, bold_label: bool = True):
            self.set_font("Helvetica", "B" if bold_label else "", 9)
            self.set_text_color(80, 80, 80)
            self.cell(55, 6, label, ln=0)
            self.set_font("Helvetica", "", 9)
            self.set_text_color(30, 30, 30)
            self.cell(0, 6, value, ln=True)
            self.set_text_color(0, 0, 0)

    pdf = ReportPDF()
    pdf.alias_nb_pages()

    # ---- Cover Page ----
    pdf.add_page()
    # Full-page gradient background using rectangles
    steps = 40
    for i in range(steps):
        r = int(30 + (37 - 30) * i / steps)
        g = int(58 + (99 - 58) * i / steps)
        b = int(95 + (235 - 95) * i / steps)
        pdf.set_fill_color(r, g, b)
        y = i * (297 / steps)
        pdf.rect(0, y, 210, 297 / steps + 1, "F")

    pdf.set_text_color(255, 255, 255)
    pdf.set_font("Helvetica", "", 10)
    pdf.set_y(60)
    pdf.cell(0, 8, "A P P S E C D", ln=True, align="C")
    pdf.ln(10)
    pdf.set_font("Helvetica", "B", 28)
    pdf.cell(0, 14, "Application Security", ln=True, align="C")
    pdf.cell(0, 14, "Assessment Report", ln=True, align="C")
    pdf.ln(6)
    pdf.set_font("Helvetica", "", 14)
    pdf.cell(0, 8, safe_text(p.get("application_name", ""), 80), ln=True, align="C")
    pdf.ln(15)

    # Cover metadata boxes
    pdf.set_font("Helvetica", "", 9)
    cover_items = [
        ("Version", p.get("application_version") or "1.0"),
        ("Environment", p.get("environment", "Staging")),
        ("Risk Level", f"{risk_level} ({risk}/100)"),
        ("Generated", gen[:10]),
    ]
    box_w = 40
    start_x = (210 - box_w * len(cover_items) - 6 * (len(cover_items) - 1)) / 2
    y_pos = pdf.get_y()
    for idx, (lbl, val) in enumerate(cover_items):
        x = start_x + idx * (box_w + 6)
        pdf.set_fill_color(255, 255, 255, )
        pdf.set_draw_color(255, 255, 255)
        # Semi-transparent box effect
        pdf.set_fill_color(255, 255, 255)
        pdf.rect(x, y_pos, box_w, 20, "D")
        pdf.set_xy(x, y_pos + 2)
        pdf.set_font("Helvetica", "", 6)
        pdf.set_text_color(200, 220, 255)
        pdf.cell(box_w, 4, lbl.upper(), align="C")
        pdf.set_xy(x, y_pos + 7)
        pdf.set_font("Helvetica", "B", 9)
        pdf.set_text_color(255, 255, 255)
        pdf.cell(box_w, 6, val[:15], align="C")

    pdf.set_y(y_pos + 30)
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(200, 220, 255)
    classification = (p.get("classification") or "Confidential").upper()
    pdf.cell(0, 8, f"--- {classification} ---", ln=True, align="C")

    pdf.set_y(250)
    pdf.set_font("Helvetica", "", 8)
    pdf.set_text_color(180, 200, 240)
    pdf.cell(0, 5, f"Target URL: {safe_text(p.get('application_url', ''), 80)}", ln=True, align="C")
    pdf.cell(0, 5, f"Testing Type: {safe_text(p.get('testing_type', ''), 30)}", ln=True, align="C")

    # ---- Executive Summary ----
    pdf.add_page()
    pdf.set_text_color(0, 0, 0)
    pdf.section_heading(1, "Executive Summary")
    ai = data.get("ai_report_content") or {}
    if ai:
        exec_sum = ai.get("executive_summary") or ai.get("ai_summary")
        if exec_sum:
            pdf.set_font("Helvetica", "B", 10)
            pdf.set_text_color(30, 58, 95)
            pdf.cell(0, 6, "AI-Generated Executive Summary", ln=True)
            pdf.set_font("Helvetica", "", 9)
            pdf.set_text_color(0, 0, 0)
            pdf.multi_cell(190, 5, safe_text(str(exec_sum), 2000))
            pdf.ln(2)
        tech = ai.get("technical_summary")
        if tech:
            pdf.set_font("Helvetica", "B", 10)
            pdf.set_text_color(30, 58, 95)
            pdf.cell(0, 6, "Technical Summary", ln=True)
            pdf.set_font("Helvetica", "", 9)
            pdf.set_text_color(0, 0, 0)
            pdf.multi_cell(190, 5, safe_text(str(tech), 2000))
            pdf.ln(2)
        strat = ai.get("strategic_recommendations")
        if strat:
            pdf.set_font("Helvetica", "B", 10)
            pdf.set_text_color(14, 165, 233)
            pdf.cell(0, 6, "Strategic Recommendations", ln=True)
            pdf.set_font("Helvetica", "", 9)
            pdf.set_text_color(0, 0, 0)
            strat_text = "\n".join(f"• {s}" for s in strat) if isinstance(strat, list) else str(strat)
            pdf.multi_cell(190, 5, safe_text(strat_text, 2000))
            pdf.ln(4)
    pdf.set_font("Helvetica", "", 10)
    pdf.multi_cell(190, 5, f"This report presents the findings of a security assessment on {safe_text(p.get('application_name', ''), 80)}. "
        "The assessment evaluated the application against OWASP Top 10, CWE Top 25, MITRE ATT&CK, ISO 27001, and NIST 800-53 frameworks.")
    pdf.ln(4)

    # Risk gauge visualization
    pdf.set_font("Helvetica", "B", 10)
    pdf.cell(0, 6, "Overall Risk Score:", ln=True)
    pdf.ln(1)
    pdf.risk_gauge(risk)
    pdf.ln(12)

    # KPI row
    pdf.set_font("Helvetica", "B", 10)
    kpi_data = [
        ("Risk Level", risk_level, sev_color_map.get(risk_level.lower(), (100, 100, 100))),
        ("Total Findings", str(len(findings)), primary_light),
        ("Test Coverage", f"{cov}%", accent_color),
        ("Critical/High", str(sev.get("critical", 0) + sev.get("high", 0)), sev_color_map["critical"]),
    ]
    kpi_w = 45
    kpi_x = 10
    kpi_y = pdf.get_y()
    for lbl, val, color in kpi_data:
        pdf.set_xy(kpi_x, kpi_y)
        pdf.set_draw_color(*primary_color)
        pdf.set_fill_color(247, 250, 252)
        pdf.rect(kpi_x, kpi_y, kpi_w, 18, "FD")
        # Top accent
        pdf.set_fill_color(*color)
        pdf.rect(kpi_x, kpi_y, kpi_w, 2.5, "F")
        pdf.set_xy(kpi_x, kpi_y + 4)
        pdf.set_font("Helvetica", "B", 14)
        pdf.set_text_color(*color)
        pdf.cell(kpi_w, 6, val, align="C")
        pdf.set_xy(kpi_x, kpi_y + 11)
        pdf.set_font("Helvetica", "", 7)
        pdf.set_text_color(100, 100, 100)
        pdf.cell(kpi_w, 4, lbl.upper(), align="C")
        kpi_x += kpi_w + 3
    pdf.set_text_color(0, 0, 0)
    pdf.set_y(kpi_y + 24)
    pdf.ln(2)

    pdf.key_value_row("Test Cases:", f"{p.get('tested_count', 0)}/{p.get('total_test_cases', 0)} executed")
    pdf.key_value_row("Target URL:", safe_text(p.get("application_url", ""), 100))
    pdf.key_value_row("Completion:", safe_text(p.get("target_completion_date") or "N/A", 30))
    pdf.ln(3)

    pdf.set_font("Helvetica", "B", 9)
    pdf.set_text_color(*primary_color)
    pdf.cell(0, 6, "Recommendation:", ln=True)
    pdf.set_text_color(0, 0, 0)
    pdf.set_font("Helvetica", "", 9)
    pdf.multi_cell(190, 5, "Address Critical and High severity findings within 30 days. "
        "Medium and Low findings should be remediated in subsequent release cycles.")
    pdf.ln(2)

    # ---- Scope & Methodology ----
    pdf.section_heading(2, "Scope & Methodology")
    scope_text = safe_text(p.get("testing_scope") or "In-scope: Application under test. Out-of-scope: Third-party systems.", 300)
    meth_type = "Black-box" if p.get("testing_type") == "black_box" else "Grey-box" if p.get("testing_type") == "grey_box" else "White-box"
    pdf.key_value_row("Scope:", scope_text)
    pdf.key_value_row("Methodology:", f"{meth_type} testing in {p.get('environment', 'staging')} environment")
    pdf.key_value_row("Standard:", "OWASP Testing Guide v4.2")
    pdf.ln(4)

    # ---- Severity Distribution ----
    pdf.section_heading(3, "Severity Distribution")
    # Table header
    pdf.set_fill_color(*primary_color)
    pdf.set_text_color(255, 255, 255)
    pdf.set_font("Helvetica", "B", 9)
    pdf.cell(80, 7, "  SEVERITY", fill=True)
    pdf.cell(30, 7, "COUNT", fill=True, align="C")
    pdf.cell(80, 7, "VISUAL", fill=True, align="C")
    pdf.ln()
    pdf.set_text_color(0, 0, 0)
    max_sev_count = max(sev.values()) if sev else 1
    alt = False
    for k, v in sorted(sev.items(), key=lambda x: _severity_order(x[0])):
        if alt:
            pdf.set_fill_color(247, 250, 252)
            pdf.rect(10, pdf.get_y(), 190, 7, "F")
        alt = not alt
        color = sev_color_map.get(k.lower(), (100, 100, 100))
        pdf.set_font("Helvetica", "", 9)
        pdf.cell(80, 7, f"  {k.capitalize()}")
        pdf.set_font("Helvetica", "B", 9)
        pdf.cell(30, 7, str(v), align="C")
        # Bar
        bar_x = pdf.get_x() + 2
        bar_y = pdf.get_y() + 1.5
        bar_max = 74
        bar_fill = max(4, (v / max_sev_count) * bar_max) if max_sev_count > 0 else 4
        pdf.set_fill_color(226, 232, 240)
        pdf.rect(bar_x, bar_y, bar_max, 4, "F")
        pdf.set_fill_color(*color)
        pdf.rect(bar_x, bar_y, bar_fill, 4, "F")
        pdf.ln()
    pdf.ln(4)

    # ---- OWASP Mapping ----
    pdf.section_heading(4, "OWASP Top 10 Mapping")
    pdf.set_fill_color(*primary_color)
    pdf.set_text_color(255, 255, 255)
    pdf.set_font("Helvetica", "B", 9)
    pdf.cell(140, 7, "  CATEGORY", fill=True)
    pdf.cell(50, 7, "COUNT", fill=True, align="C")
    pdf.ln()
    pdf.set_text_color(0, 0, 0)
    alt = False
    for k, v in sorted(owasp.items()):
        if alt:
            pdf.set_fill_color(247, 250, 252)
            pdf.rect(10, pdf.get_y(), 190, 7, "F")
        alt = not alt
        pdf.set_font("Helvetica", "", 9)
        pdf.cell(140, 7, f"  {safe_text(k, 60)}")
        pdf.set_font("Helvetica", "B", 9)
        pdf.cell(50, 7, str(v), align="C")
        pdf.ln()
    pdf.ln(4)

    # ---- CWE Mapping ----
    pdf.section_heading(5, "CWE Mapping")
    pdf.set_fill_color(*primary_color)
    pdf.set_text_color(255, 255, 255)
    pdf.set_font("Helvetica", "B", 9)
    pdf.cell(50, 7, "  CWE ID", fill=True)
    pdf.cell(90, 7, "  FINDING", fill=True)
    pdf.cell(50, 7, "SEVERITY", fill=True, align="C")
    pdf.ln()
    pdf.set_text_color(0, 0, 0)
    alt = False
    for f in findings:
        if alt:
            pdf.set_fill_color(247, 250, 252)
            pdf.rect(10, pdf.get_y(), 190, 7, "F")
        alt = not alt
        pdf.set_font("Helvetica", "", 8)
        pdf.cell(50, 7, f"  {safe_text(f.get('cwe_id') or '-', 20)}")
        pdf.cell(90, 7, f"  {safe_text(f.get('title', ''), 45)}")
        x_before = pdf.get_x()
        y_before = pdf.get_y()
        pdf.set_xy(x_before, y_before + 1)
        pdf.severity_badge(f.get("severity", ""))
        pdf.set_xy(x_before + 50, y_before)
        pdf.ln()
    pdf.ln(4)

    # ---- Compliance Frameworks ----
    pdf.section_heading(6, "Compliance Frameworks")
    pdf.set_font("Helvetica", "", 9)
    pdf.multi_cell(190, 5, "All findings are mapped to the following compliance frameworks:")
    pdf.ln(2)
    frameworks = [
        ("OWASP Top 10 (2021)", "Web Application Security Standard"),
        ("CWE Top 25", "Common Weakness Enumeration"),
        ("MITRE ATT&CK", "Adversarial Tactics and Techniques"),
        ("ISO 27001 Annex A", "Information Security Controls"),
        ("NIST 800-53", "Security and Privacy Controls"),
    ]
    pdf.set_fill_color(*primary_color)
    pdf.set_text_color(255, 255, 255)
    pdf.set_font("Helvetica", "B", 9)
    pdf.cell(70, 7, "  FRAMEWORK", fill=True)
    pdf.cell(90, 7, "  DESCRIPTION", fill=True)
    pdf.cell(30, 7, "STATUS", fill=True, align="C")
    pdf.ln()
    pdf.set_text_color(0, 0, 0)
    alt = False
    for fw_name, fw_desc in frameworks:
        if alt:
            pdf.set_fill_color(247, 250, 252)
            pdf.rect(10, pdf.get_y(), 190, 7, "F")
        alt = not alt
        pdf.set_font("Helvetica", "B", 8)
        pdf.cell(70, 7, f"  {fw_name}")
        pdf.set_font("Helvetica", "", 8)
        pdf.cell(90, 7, f"  {fw_desc}")
        pdf.set_fill_color(22, 163, 106)
        pdf.set_text_color(255, 255, 255)
        pdf.set_font("Helvetica", "B", 7)
        x_b = pdf.get_x()
        pdf.cell(30, 7, "MAPPED", align="C", fill=True)
        pdf.set_text_color(0, 0, 0)
        pdf.ln()
    pdf.ln(4)

    # ---- Detailed Findings ----
    pdf.section_heading(7, "Detailed Findings with Evidence")
    w = 190
    for i, f in enumerate(findings, 1):
        pdf.add_page()
        sev_lower = f.get("severity", "").lower()
        color = sev_color_map.get(sev_lower, (100, 100, 100))

        # Finding header with severity color accent
        pdf.set_fill_color(*color)
        pdf.rect(10, pdf.get_y(), 190, 1.5, "F")
        pdf.ln(3)

        pdf.set_font("Helvetica", "B", 12)
        pdf.set_text_color(*primary_color)
        pdf.cell(0, 7, f"Finding #{i}: {safe_text(f.get('title',''), 65)}", ln=True)
        pdf.set_text_color(0, 0, 0)

        # Severity badge
        x_badge = pdf.get_x()
        pdf.severity_badge(f.get("severity", ""))
        if f.get("cvss_score"):
            pdf.set_font("Helvetica", "", 9)
            pdf.cell(5, 5, "")
            pdf.cell(0, 5, f"CVSS: {f.get('cvss_score')}", ln=True)
        else:
            pdf.ln()
        pdf.ln(2)

        # Meta grid
        pdf.set_fill_color(247, 250, 252)
        meta_y = pdf.get_y()
        pdf.rect(10, meta_y, 190, 18, "F")
        pdf.set_xy(12, meta_y + 1)
        pdf.set_font("Helvetica", "", 7)
        pdf.set_text_color(100, 100, 100)
        pdf.cell(45, 4, "AFFECTED URL")
        pdf.cell(45, 4, "OWASP CATEGORY")
        pdf.cell(35, 4, "CWE ID")
        pdf.cell(55, 4, "CVSS SCORE")
        pdf.ln()
        pdf.set_xy(12, meta_y + 6)
        pdf.set_font("Helvetica", "B", 8)
        pdf.set_text_color(30, 30, 30)
        pdf.cell(45, 5, safe_text(f.get("affected_url") or "-", 25))
        pdf.cell(45, 5, safe_text(f.get("owasp_category") or "-", 25))
        pdf.cell(35, 5, safe_text(f.get("cwe_id") or "-", 15))
        pdf.cell(55, 5, safe_text(f.get("cvss_score") or "N/A", 10))
        pdf.set_text_color(0, 0, 0)
        pdf.set_y(meta_y + 20)

        # Description
        pdf.set_font("Helvetica", "B", 9)
        pdf.set_text_color(*primary_color)
        pdf.cell(0, 6, "Description", ln=True)
        pdf.set_text_color(0, 0, 0)
        pdf.set_font("Helvetica", "", 9)
        pdf.multi_cell(w, 5, safe_text(f.get("description"), 500))
        pdf.ln(2)

        # Impact
        pdf.set_font("Helvetica", "B", 9)
        pdf.set_text_color(*primary_color)
        pdf.cell(0, 6, "Impact", ln=True)
        pdf.set_text_color(0, 0, 0)
        pdf.set_font("Helvetica", "", 9)
        pdf.multi_cell(w, 5, safe_text(f.get("impact"), 400))
        pdf.ln(2)

        # Reproduction Steps
        pdf.set_font("Helvetica", "B", 9)
        pdf.set_text_color(*primary_color)
        pdf.cell(0, 6, "Reproduction Steps", ln=True)
        pdf.set_text_color(0, 0, 0)
        pdf.set_font("Helvetica", "", 9)
        pdf.multi_cell(w, 5, safe_text(f.get("reproduction_steps"), 400))
        pdf.ln(2)

        # Evidence
        evidence_list = f.get("evidence") or []
        for ev_idx, ev in enumerate(evidence_list):
            if isinstance(ev, dict):
                fpath = _get_evidence_file_path(project_id, ev)
                if fpath and fpath.exists() and fpath.suffix.lower() in {".png", ".jpg", ".jpeg", ".gif", ".webp"}:
                    try:
                        pdf.set_fill_color(*primary_color)
                        pdf.set_text_color(255, 255, 255)
                        pdf.set_font("Helvetica", "B", 8)
                        fn = ev.get("filename", "Screenshot")
                        pdf.cell(0, 6, f"  Evidence {ev_idx + 1}: {fn}", fill=True, ln=True)
                        pdf.set_text_color(0, 0, 0)
                        pdf.ln(1)
                        pdf.image(str(fpath), x=10, w=180)
                        pdf.ln(3)
                    except Exception:
                        pass

        # Request/Response (Automated Evidence)
        if f.get("request"):
            pdf.set_font("Helvetica", "B", 9)
            pdf.set_text_color(*primary_color)
            pdf.cell(0, 6, "Request (Automated Evidence)", ln=True)
            pdf.set_text_color(0, 0, 0)
            pdf.set_font("Helvetica", "", 8)
            pdf.multi_cell(w, 5, safe_text(f.get("request"), 600))
            pdf.ln(2)
        if f.get("response"):
            pdf.set_font("Helvetica", "B", 9)
            pdf.set_text_color(*primary_color)
            pdf.cell(0, 6, "Response (Automated Evidence)", ln=True)
            pdf.set_text_color(0, 0, 0)
            pdf.set_font("Helvetica", "", 8)
            pdf.multi_cell(w, 5, safe_text(f.get("response"), 600))
            pdf.ln(2)

        # Recommendation
        pdf.set_font("Helvetica", "B", 9)
        pdf.set_text_color(*primary_color)
        pdf.cell(0, 6, "Recommendation", ln=True)
        pdf.set_text_color(0, 0, 0)
        pdf.set_font("Helvetica", "", 9)
        pdf.multi_cell(w, 5, safe_text(f.get("recommendation"), 400))
        pdf.ln(2)

        # Compliance mapping
        comp = f.get("compliance") or {}
        comp_parts = []
        if comp.get("mitre_attack"):
            comp_parts.append(f"ATT&CK: {comp['mitre_attack']}")
        if comp.get("iso_27001"):
            comp_parts.append(f"ISO: {comp['iso_27001']}")
        if comp.get("nist_800_53"):
            comp_parts.append(f"NIST: {comp['nist_800_53']}")
        if comp_parts:
            pdf.set_font("Helvetica", "B", 8)
            pdf.set_text_color(100, 100, 100)
            pdf.cell(0, 5, "Compliance: " + " | ".join(comp_parts), ln=True)
            pdf.set_text_color(0, 0, 0)

    # ---- Final page: Confidentiality ----
    pdf.add_page()
    pdf.ln(30)
    pdf.set_fill_color(*primary_color)
    pdf.rect(10, pdf.get_y(), 190, 50, "F")
    y_box = pdf.get_y()
    pdf.set_xy(15, y_box + 5)
    pdf.set_text_color(255, 255, 255)
    pdf.set_font("Helvetica", "B", 11)
    pdf.cell(0, 7, "CONFIDENTIALITY NOTICE", ln=True)
    pdf.set_x(15)
    pdf.set_font("Helvetica", "", 8)
    pdf.multi_cell(180, 4,
        "This document contains confidential and privileged information intended solely for the authorized "
        "recipient(s). Any unauthorized review, use, disclosure, or distribution is strictly prohibited. "
        "If you have received this document in error, please notify the sender immediately and destroy all copies. "
        "This report and its contents are protected under applicable confidentiality agreements and information security policies.")
    pdf.set_text_color(0, 0, 0)

    buf = BytesIO()
    pdf.output(buf)
    buf.seek(0)
    return buf.getvalue()


def generate_json(data: dict) -> str:
    import json
    out = dict(data)
    for f in out.get("findings", []):
        f.pop("compliance", None)
    return json.dumps(out, indent=2)


def generate_csv(data: dict) -> str:
    import csv
    from io import StringIO
    findings = data["findings"]
    if not findings:
        return "title,severity,owasp_category,cwe_id,affected_url,description,impact,recommendation\n"
    out = StringIO()
    keys = ["title", "severity", "owasp_category", "cwe_id", "affected_url", "description", "impact", "recommendation"]
    w = csv.DictWriter(out, fieldnames=keys, extrasaction="ignore")
    w.writeheader()
    for f in findings:
        w.writerow({k: (f.get(k) or "") for k in keys})
    return out.getvalue()
